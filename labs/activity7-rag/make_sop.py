#!/usr/bin/env python3
"""Generate a sample company HR SOP Word document for RAG upload demo."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def add_section(doc, number, title, intro, clauses):
    doc.add_heading(f"{number}. {title}", level=1)
    add_para(doc, intro)
    for ctitle, ctext in clauses:
        doc.add_heading(ctitle, level=2)
        if isinstance(ctext, list):
            add_bullets(doc, ctext)
        else:
            add_para(doc, ctext)
    doc.add_paragraph()


doc = Document()

# ---- Cover ----
title = doc.add_heading("MyCompany Singapore", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("Human Resources — Standard Operating Procedures (SOP)")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(14)

meta = doc.add_paragraph(
    "Document Owner: Human Resources Department\n"
    "Version: 1.0    |    Effective Date: 1 January 2026    |    Classification: Internal\n"
    "Applies to: All permanent, contract, and probationary employees of MyCompany Singapore"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in meta.runs:
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
purpose = doc.add_paragraph()
purpose.add_run("Purpose. ").bold = True
purpose.add_run(
    "This handbook consolidates the company's core people policies into a single reference. "
    "It is the authoritative source for the standard operating procedures listed below. "
    "Employees should consult this document first; for anything not covered here, contact "
    "hr@mycompany-sg.example.com."
)

doc.add_heading("Contents", level=1)
add_bullets(doc, [
    "1. Leave Policy",
    "2. Medical Certificate (MC) Policy",
    "3. Performance Management Policy",
    "4. PDPA (Personal Data Protection Act) Policy",
    "5. Marketing Policy",
    "6. Public Relations Policy",
    "7. Data Privacy Policy",
])
doc.add_page_break()

# ---- 1. Leave Policy ----
add_section(
    doc, "1", "Leave Policy",
    "This policy governs all forms of paid and unpaid leave. Leave must be applied for "
    "through the HR portal and approved by the reporting manager before it is taken, except "
    "in genuine emergencies.",
    [
        ("1.1 Annual Leave", [
            "Employees are entitled to 14 days of paid annual leave per calendar year, increasing by 1 day for every completed year of service, up to a maximum of 21 days.",
            "A minimum of 3 working days' notice is required for leave of 1–2 days; 2 weeks' notice for 3 or more consecutive days.",
            "Up to 5 unused annual leave days may be carried forward to the next year and must be cleared by 31 March.",
        ]),
        ("1.2 Sick Leave", [
            "Employees are entitled to 14 days of paid outpatient sick leave and 60 days of paid hospitalisation leave per year, inclusive of outpatient days.",
            "Sick leave of more than one day must be supported by a valid Medical Certificate (see Section 2).",
        ]),
        ("1.3 Other Leave", [
            "Maternity leave: 16 weeks of paid leave for eligible employees.",
            "Paternity leave: 2 weeks of paid leave.",
            "Childcare leave: 6 days per year for parents of children under 7.",
            "Compassionate leave: up to 3 days for the death of an immediate family member.",
            "Unpaid leave may be granted at management's discretion once paid entitlements are exhausted.",
        ]),
        ("1.4 Application Procedure",
         "All leave is applied for via the HR self-service portal. The reporting manager approves "
         "or rejects within 2 working days. Approved leave is automatically reflected in the team "
         "calendar. Emergency leave should be reported to the manager by phone or message before "
         "9:00 AM on the day, and formalised in the portal within 24 hours."),
    ],
)

# ---- 2. MC Policy ----
add_section(
    doc, "2", "Medical Certificate (MC) Policy",
    "A Medical Certificate (MC) is the official proof of an employee's medical unfitness for work. "
    "This policy explains when an MC is required and how it must be submitted.",
    [
        ("2.1 When an MC Is Required", [
            "An MC is required for any sick leave of more than one (1) consecutive day.",
            "An MC may be requested for a single day of sick leave if the employee has a recurring pattern of single-day absences.",
            "MCs must be issued by a registered medical practitioner or a government-approved clinic or hospital.",
        ]),
        ("2.2 Submission", [
            "Upload a clear photo or scan of the MC to the HR portal within 48 hours of returning to work.",
            "The MC must show the employee's name, the clinic's stamp, the doctor's signature, the date of consultation, and the number of days certified.",
            "Backdated MCs are not accepted unless the doctor has formally certified the backdated period.",
        ]),
        ("2.3 Misuse",
         "Submitting a forged or altered MC is a serious offence and may result in disciplinary action "
         "up to and including termination. HR reserves the right to verify any MC directly with the "
         "issuing clinic."),
    ],
)

# ---- 3. Performance Policy ----
add_section(
    doc, "3", "Performance Management Policy",
    "This policy describes how employee performance is set, reviewed, and rewarded to ensure fair, "
    "transparent, and growth-oriented evaluation.",
    [
        ("3.1 Goal Setting",
         "At the start of each performance year, every employee agrees on SMART objectives with their "
         "manager. Objectives are documented in the performance system and reviewed quarterly."),
        ("3.2 Review Cycle", [
            "Mid-year review (July): an informal check-in on progress and obstacles.",
            "Year-end review (December): a formal appraisal rating overall performance.",
            "Ratings use a 5-point scale: 1 = Below Expectations, 3 = Meets Expectations, 5 = Exceptional.",
        ]),
        ("3.3 Outcomes", [
            "Annual increments and bonuses are linked to the year-end rating and company performance.",
            "Employees rated Below Expectations are placed on a 90-day Performance Improvement Plan (PIP) with defined milestones and support.",
            "Promotion is based on sustained performance, demonstrated competencies, and business need.",
        ]),
        ("3.4 Appeals",
         "An employee who disagrees with a rating may raise an appeal to HR within 14 days of the "
         "review. HR will facilitate a discussion with a skip-level manager to reach a fair outcome."),
    ],
)

# ---- 4. PDPA Policy ----
add_section(
    doc, "4", "PDPA (Personal Data Protection Act) Policy",
    "This policy sets out how MyCompany Singapore complies with Singapore's Personal Data Protection "
    "Act (PDPA) when handling the personal data of employees, customers, and partners.",
    [
        ("4.1 Consent & Purpose", [
            "Personal data is collected only for legitimate business purposes that have been communicated to the individual.",
            "Consent is obtained before collecting, using, or disclosing personal data, except where the law permits otherwise.",
        ]),
        ("4.2 The Nine PDPA Obligations",
         "All staff must observe the PDPA obligations: Consent, Purpose Limitation, Notification, "
         "Access & Correction, Accuracy, Protection, Retention Limitation, Transfer Limitation, and "
         "Accountability."),
        ("4.3 Data Protection Officer (DPO)",
         "The company appoints a Data Protection Officer responsible for PDPA compliance. Questions, "
         "access requests, and complaints should be directed to dpo@mycompany-sg.example.com."),
        ("4.4 Breach Notification", [
            "Any suspected data breach must be reported to the DPO within 2 hours of discovery.",
            "Notifiable breaches will be reported to the PDPC and affected individuals within the timelines set by law (generally within 3 calendar days of assessment).",
        ]),
    ],
)

# ---- 5. Marketing Policy ----
add_section(
    doc, "5", "Marketing Policy",
    "This policy governs how the company markets its products and services to protect brand "
    "integrity and ensure honest, lawful communication.",
    [
        ("5.1 Brand & Messaging", [
            "All marketing materials must use approved logos, colours, and the current brand guidelines.",
            "Claims about products must be accurate, substantiated, and never misleading.",
            "Competitor comparisons must be factual and fair.",
        ]),
        ("5.2 Approvals",
         "All external campaigns require sign-off from the Marketing Lead. Campaigns referencing "
         "financial results, legal matters, or partnerships additionally require Legal and Finance "
         "review before release."),
        ("5.3 Digital & Email Marketing", [
            "Email marketing must comply with the PDPA's Do Not Call provisions and include an unsubscribe option.",
            "Use only opt-in marketing lists; never purchase third-party contact lists.",
        ]),
    ],
)

# ---- 6. Public Relations Policy ----
add_section(
    doc, "6", "Public Relations Policy",
    "This policy ensures the company speaks to the public, media, and on social platforms with one "
    "consistent and authorised voice.",
    [
        ("6.1 Media Enquiries",
         "Only the designated Corporate Communications spokesperson may speak to the media on behalf "
         "of the company. Employees receiving media enquiries must forward them to "
         "pr@mycompany-sg.example.com without comment."),
        ("6.2 Social Media", [
            "Official company channels are managed solely by the Communications team.",
            "Employees posting personal opinions online must not present them as the company's position and must not disclose confidential information.",
        ]),
        ("6.3 Crisis Communication",
         "In a crisis, all communication is centralised under the Crisis Communications Team. No "
         "employee may issue statements, post on social media about the incident, or speak to "
         "journalists until an official position is released."),
    ],
)

# ---- 7. Data Privacy Policy ----
add_section(
    doc, "7", "Data Privacy Policy",
    "This policy defines the day-to-day security practices that protect the confidentiality, "
    "integrity, and availability of company and personal data. It complements the PDPA Policy "
    "(Section 4) with practical controls.",
    [
        ("7.1 Access Control", [
            "Access to systems and data follows the principle of least privilege — staff receive only the access needed for their role.",
            "Multi-factor authentication is mandatory for all corporate accounts.",
            "Passwords must be at least 12 characters and are never shared.",
        ]),
        ("7.2 Handling & Storage", [
            "Confidential data must be stored on approved company systems, never on personal devices or unapproved cloud services.",
            "Sensitive data must be encrypted in transit and at rest.",
            "Printed confidential documents must be shredded when no longer needed.",
        ]),
        ("7.3 Retention & Disposal",
         "Personal data is retained only as long as necessary for the stated purpose or as required "
         "by law, after which it is securely deleted or anonymised."),
        ("7.4 Employee Responsibilities", [
            "Lock your screen when away from your desk.",
            "Report lost devices or suspected security incidents to IT Security immediately.",
            "Complete the annual data privacy and security awareness training.",
        ]),
    ],
)

doc.add_paragraph()
end = doc.add_paragraph(
    "End of document. This SOP is reviewed annually by the Human Resources Department. "
    "For clarification on any policy, contact hr@mycompany-sg.example.com."
)
for r in end.runs:
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

out = "MyCompany-HR-SOP.docx"
doc.save(out)
print(f"Saved {out}")
