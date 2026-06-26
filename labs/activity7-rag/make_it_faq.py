#!/usr/bin/env python3
"""Generate a sample IT Support FAQ Word document for the Activity 5 RAG upload demo."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_steps(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")


def add_bullets(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def faq(doc, number, question, intro, steps=None, bullets=None, note=None):
    doc.add_heading(f"FAQ {number}. {question}", level=1)
    if intro:
        add_para(doc, intro)
    if steps:
        doc.add_heading("Steps", level=2)
        add_steps(doc, steps)
    if bullets:
        add_bullets(doc, bullets)
    if note:
        p = doc.add_paragraph()
        r = p.add_run("Note: " + note)
        r.italic = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()


doc = Document()

# ---- Cover ----
title = doc.add_heading("MyCompany Singapore", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("IT Service Desk — Frequently Asked Questions (FAQ)")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].bold = True
sub.runs[0].font.size = Pt(14)

meta = doc.add_paragraph(
    "Document Owner: IT Service Desk\n"
    "Version: 1.0    |    Effective Date: 1 January 2026    |    Classification: Internal\n"
    "Service Desk: ithelpdesk@mycompany-sg.example.com    |    IT Portal: https://itportal.mycompany-sg.example.com"
)
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in meta.runs:
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Purpose. ").bold = True
p.add_run(
    "This FAQ is the first-line reference for common IT issues at MyCompany Singapore. "
    "Try the steps here before raising a ticket. If a problem is not covered or the steps "
    "do not resolve it, contact the IT Service Desk with your name, asset tag, the exact "
    "error message, and what you have already tried."
)

doc.add_heading("Contents", level=1)
add_bullets(doc, [
    "FAQ 1. Reset a forgotten or expired password",
    "FAQ 2. Unlock a locked account",
    "FAQ 3. Set up or fix Multi-Factor Authentication (MFA)",
    "FAQ 4. Connect to the corporate VPN",
    "FAQ 5. Wi-Fi / network connection problems",
    "FAQ 6. Email and Outlook issues",
    "FAQ 7. Install approved software",
    "FAQ 8. Printer setup and printing problems",
    "FAQ 9. Laptop / hardware troubleshooting",
    "FAQ 10. Request new equipment or access",
    "FAQ 11. Shared drive / file access",
    "FAQ 12. Report a phishing email or security incident",
    "FAQ 13. How to raise an IT support ticket",
])
doc.add_page_break()

faq(doc, "1", "How do I reset a forgotten or expired password?",
    "Use the self-service password reset (SSPR) portal — you do not need to call the Service Desk for a routine reset.",
    steps=[
        "Go to https://passwordreset.mycompany-sg.example.com from any browser.",
        "Enter your work email and complete the identity check (MFA prompt or security questions).",
        "Choose a new password that meets the policy: at least 12 characters, with upper- and lower-case letters, a number, and a symbol.",
        "Sign out and sign back in to all devices using the new password.",
    ],
    note="Passwords expire every 90 days. You will see a reminder 14 days before expiry. Never share your password or reuse it on personal sites.")

faq(doc, "2", "My account is locked — how do I unlock it?",
    "Accounts lock after 5 failed sign-in attempts and unlock automatically after 15 minutes.",
    steps=[
        "Wait 15 minutes, then try again with the correct password.",
        "If you have forgotten the password, run a self-service reset (see FAQ 1) — this also clears the lock.",
        "If it is still locked after a reset, raise a ticket with the Service Desk and quote 'account locked'.",
    ])

faq(doc, "3", "How do I set up or fix Multi-Factor Authentication (MFA)?",
    "MFA (the Authenticator app) is required for all corporate accounts.",
    steps=[
        "Install 'Microsoft Authenticator' from the App Store or Google Play.",
        "On a computer, go to https://aka.ms/mfasetup and sign in.",
        "Choose 'Add sign-in method' → 'Authenticator app', then scan the QR code with the app.",
        "Approve the test notification to finish enrolment.",
    ],
    bullets=[
        "Lost or replaced your phone? Raise a ticket to have MFA reset, then re-enrol.",
        "Not receiving prompts? Check the phone has internet and notifications are enabled for the Authenticator app.",
    ])

faq(doc, "4", "How do I connect to the corporate VPN?",
    "Use GlobalConnect VPN to access internal systems when working off-site.",
    steps=[
        "Open the 'GlobalConnect' app (pre-installed on company laptops; otherwise install it from the IT Portal).",
        "Enter the gateway address: vpn.mycompany-sg.example.com.",
        "Sign in with your work email and password, then approve the MFA prompt.",
        "Wait for the status to show 'Connected' before opening internal sites.",
    ],
    note="If the VPN keeps dropping, switch from Wi-Fi to a wired connection or move closer to your router, then reconnect. Corporate VPN is required for HR and finance systems.")

faq(doc, "5", "I can't connect to Wi-Fi or the network — what do I do?",
    "Most connection problems are fixed by reconnecting to the right network.",
    steps=[
        "Confirm you are joined to 'QCOM-Corp' (staff), not 'QCOM-Guest'.",
        "Toggle Wi-Fi off and on, or use Airplane mode for 10 seconds.",
        "Forget the 'QCOM-Corp' network and rejoin, entering your work credentials when prompted.",
        "Restart the laptop if the issue persists.",
    ],
    note="If no networks appear at all, the wireless adapter may be disabled — raise a hardware ticket.")

faq(doc, "6", "My email / Outlook is not working — how do I fix it?",
    "Common Outlook issues (not sending/receiving, stuck mail) usually clear with these steps.",
    steps=[
        "Check your internet/VPN connection first.",
        "In Outlook, look at the bottom status bar — if it says 'Disconnected', click 'Send/Receive' → 'Work Offline' to toggle back online.",
        "Empty the Outbox if a large attachment is stuck (over 25 MB will not send — use the file-share link instead).",
        "Restart Outlook; if it still fails, restart the laptop.",
    ],
    bullets=[
        "Mailbox full? Archive old items or empty Deleted Items — the limit is 50 GB.",
        "Webmail alternative: https://mail.mycompany-sg.example.com works from any browser.",
    ])

faq(doc, "7", "How do I install approved software?",
    "Approved applications are available from the Company Portal — no admin rights needed.",
    steps=[
        "Open the 'Company Portal' app on your laptop.",
        "Search for the application and click 'Install'.",
        "Wait for the install to finish (the status changes to 'Installed').",
    ],
    note="Software not listed in the Company Portal requires manager approval. Raise a software request ticket with the business justification. Do not install software from the internet — it is blocked by policy.")

faq(doc, "8", "How do I set up a printer or fix printing problems?",
    "Office printers use follow-me printing — your job is released at any printer when you tap your badge.",
    steps=[
        "Open 'Company Portal' → 'Printers' and add 'QCOM-FollowMe'.",
        "Print as usual, then tap your staff badge on any enabled printer to release the job.",
        "If nothing prints, check the printer screen for paper/toner alerts and try another printer.",
    ],
    note="Jobs not collected within 24 hours are automatically deleted. For a paper jam or hardware fault, log a ticket with the printer's location/ID.")

faq(doc, "9", "My laptop is slow, frozen, or won't start — what should I try?",
    "Work through these basic checks before logging a hardware ticket.",
    steps=[
        "Save your work and restart the laptop — this resolves most slowness and freezes.",
        "Confirm it is charged and the power adapter is firmly connected.",
        "Close unused applications and browser tabs to free memory.",
        "Run pending updates from 'Company Portal' → 'Updates'.",
    ],
    note="If the laptop will not power on, shows a blue/black screen, or makes unusual noises, raise an urgent hardware ticket and include the asset tag from the sticker on the base.")

faq(doc, "10", "How do I request new equipment or extra access?",
    "All new hardware, peripherals, and system access go through a request ticket for approval.",
    bullets=[
        "Hardware (laptop, monitor, headset, dock): raise an 'Equipment Request' with your manager's name for approval.",
        "System / application access (e.g. a shared mailbox, a business system): raise an 'Access Request' naming the system and the reason.",
        "Standard requests are fulfilled within 3–5 business days; approvals from your manager speed this up.",
    ])

faq(doc, "11", "I can't open a shared drive or file — how do I get access?",
    "Shared drives are permission-controlled by team.",
    steps=[
        "Make sure you are connected to the VPN (see FAQ 4) — shared drives are not reachable off the corporate network.",
        "Re-map the drive: File Explorer → 'This PC' → 'Map network drive' → enter the path provided by your team.",
        "If you get 'Access denied', raise an 'Access Request' ticket naming the exact folder and your manager for approval.",
    ])

faq(doc, "12", "How do I report a phishing email or a security incident?",
    "Report suspicious messages immediately — do not click links or open attachments.",
    steps=[
        "In Outlook, select the suspicious email and click the 'Report Phishing' button on the toolbar.",
        "If there is no button, forward the email to phishing@mycompany-sg.example.com.",
        "Delete the email after reporting.",
        "If you already clicked a link or entered your password, change your password now (FAQ 1) and call the Service Desk immediately.",
    ],
    note="Report lost or stolen devices to the Service Desk at once so the device can be remotely locked and wiped.")

faq(doc, "13", "How do I raise an IT support ticket?",
    "When the FAQ does not resolve your issue, log a ticket so the Service Desk can help.",
    bullets=[
        "Portal: https://itportal.mycompany-sg.example.com → 'New Request'.",
        "Email: ithelpdesk@mycompany-sg.example.com (a ticket is created automatically).",
        "Phone (urgent / system down): +65 6000 0000, Mon–Fri 8:30am–6:00pm.",
    ])
add_para(doc,
    "Always include: your full name, asset tag, the exact error message, when it started, and the steps you "
    "have already tried. Priority: P1 (business-wide outage) within 1 hour; P2 (you cannot work) within 4 hours; "
    "P3 (general request) within 1–2 business days.")

doc.add_paragraph()
end = doc.add_paragraph(
    "End of document. This FAQ is reviewed quarterly by the IT Service Desk. "
    "For anything not covered here, contact ithelpdesk@mycompany-sg.example.com."
)
for r in end.runs:
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

out = "MyCompany-IT-Support-FAQ.docx"
doc.save(out)
print(f"Saved {out}")
