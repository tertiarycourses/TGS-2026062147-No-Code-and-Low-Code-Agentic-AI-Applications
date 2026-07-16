#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'No Code and Low Code Agentic AI Applications' (TGS-2026062147):
  - Written Assessment (SAQ)  — 12 open-ended KNOWLEDGE questions (K1–K12), aligned to the slides
  - Practical Performance (PP) — 7 PRACTICAL tasks (LO1–LO7), aligned to the in-class activities
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide) + version control.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "No Code and Low Code Agentic AI Applications"   # <<Course Title>>
COURSE_CODE = "TGS-2026062147"                                  # <<Course Code, e.g. TGS-XXXXXXXXXX>>
# ────────────────────────────────────────────────────────────────────────────
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own .claude/skills/tertiary-course-slides/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, ".claude/skills/tertiary-course-slides/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("n8n-course-logo.png")

Q_VER, A_VER = "v1", "v1"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points])
WRITTEN = [
 ("K1",
  "Generative AI systems are typically designed to produce content in response to user prompts, while agentic "
  "AI systems are built to plan, decide, and take actions autonomously to achieve goals over multiple steps.",
  "What are the key differences between Generative AI and Agentic AI in terms of capabilities, behaviour and use cases?",
  ["Autonomy: Generative AI responds to a prompt; Agentic AI plans, decides and acts on its own.",
   "Goals: Generative AI handles a single task; Agentic AI pursues a multi-step goal.",
   "Tools/memory: Agentic AI calls tools/APIs, keeps memory and looks up data (RAG); Generative AI does not.",
   "Use cases: content generation vs. autonomous, multi-step workflows (e.g. a Telegram support agent)."]),
 ("K2",
  "In n8n, workflows can be started in different ways depending on how and when automation should occur — by a "
  "user action, a time-based event, or a request from an external system.",
  "What are the key trigger nodes in n8n that can be used to start a workflow?",
  ["Manual Trigger", "Form Trigger (form submission)", "Schedule Trigger (time-based)",
   "Webhook (external HTTP request)", "Telegram Trigger / Chat Trigger (incoming message)"]),
 ("K3",
  "After a workflow is triggered, action nodes process data, apply logic and control the flow — enabling "
  "decision-making, branching, data handling and custom logic.",
  "What are some key action / logic nodes in n8n used to control workflow logic and data processing?",
  ["IF node (branching on a condition)", "Switch node (multi-way routing)", "Merge node",
   "Edit Fields (Set) node", "Code node (JavaScript/Python)", "Filter node"]),
 ("K4",
  "An AI agent perceives context, reasons about a task and takes actions autonomously to achieve a goal. To work, "
  "it must combine reasoning with memory and the ability to use external systems.",
  "What are the key components required to build a functional AI agent in n8n?",
  ["A Chat Model (LLM) for reasoning — e.g. OpenAI or Google Gemini.",
   "Memory (e.g. a window buffer) to retain conversation context.",
   "Tools (Vector Store, Data Table, HTTP, sub-workflows) the agent can call.",
   "A System Instruction that defines the agent's scope, rules and how to route between tools."]),
 ("K5",
  "You are building an AI system that must retrieve relevant information from your own documents and use that "
  "context to generate accurate, up-to-date answers with an LLM — improving factual grounding and reducing "
  "hallucination without retraining the model.",
  "Which type of AI workflow best supports this combination of information retrieval and language generation, and how does it work?",
  ["Retrieval-Augmented Generation (RAG).",
   "Documents are split into chunks, embedded, and stored in a vector store.",
   "At query time the question is embedded and the closest chunks are retrieved.",
   "The retrieved chunks are given to the LLM as context so it answers only from your documents."]),
 ("K6",
  "You are designing an n8n workflow that must fetch data from external systems and online services (e.g. a "
  "market-data or news API) and use the returned data in later steps.",
  "Which method would you use in n8n to connect to external data sources and services, and what do you configure on it?",
  ["The HTTP Request node (or a built-in service node).",
   "Configure the method (GET/POST), URL, headers and query parameters.",
   "Store API keys in credentials (e.g. Query Auth / header auth) — never hard-coded.",
   "Use the structured JSON response in downstream nodes."]),
 ("K7",
  "In a RAG system, text must be turned into a form that supports meaning-based (semantic) search before it can "
  "be stored and retrieved. The vector store must also be configured to match the embedding model.",
  "Explain how embeddings and a vector store enable semantic retrieval in a RAG pipeline, and state the "
  "dimension rule that makes ingestion succeed.",
  ["An embedding turns a chunk of text into a vector (a list of numbers capturing meaning).",
   "Similar meanings produce vectors that are close together.",
   "Vectors are saved in a vector store; the query is embedded and the nearest vectors are retrieved (e.g. cosine similarity).",
   "The embedding model's output dimension must equal the store's configured dimension "
   "(e.g. OpenAI text-embedding-3-small = 1536; Gemini gemini-embedding-001 = 3072)."]),
 ("K8",
  "An in-memory vector store is simple but is lost when the workflow restarts. For production RAG you want a "
  "store that persists and scales.",
  "Name persistent vector databases you can use for RAG in n8n, and one distinguishing trait of each.",
  ["Pinecone — fully-managed SaaS; just create an index (dimension + metric), zero-ops.",
   "Supabase (pgvector) — Postgres + a vector extension; good if you already use Postgres.",
   "Qdrant — open-source; run via Docker or Qdrant Cloud for full control.",
   "(All must match the embedding dimension, e.g. 1536 / 3072.)"]),
 ("K9",
  "Some workflows must start automatically the moment an event happens outside n8n — a form submission, a website "
  "chat message, or a call from another application — rather than on a schedule or manual run.",
  "What is the primary purpose of a Webhook in n8n, and what node pairs with it to reply to the caller?",
  ["A Webhook is a URL that external systems call to trigger the workflow in real time.",
   "Use cases: website chat, form submissions, payment/GitHub/Stripe events.",
   "Pair it with a Respond to Webhook node to send a reply back to the caller.",
   "Set Allowed Origins (CORS) = * so a browser page can call it."]),
 ("K10",
  "To make an AI agent trustworthy, you place safety checks around it so that unsafe input never reaches the "
  "agent and unsafe output never reaches the user.",
  "What is the difference between an input (pre) guardrail and an output (post) guardrail, and what does each block?",
  ["Input / pre-guardrail: an LLM classifies each incoming message ALLOW or BLOCK before it reaches the agent.",
   "It blocks prompt injection, jailbreaks and requests for another person's private data.",
   "Output / post-guardrail: an LLM checks every reply SAFE or LEAK before it reaches the user.",
   "It blocks salary figures, NRIC, credentials or system-instruction leaks; a blocked case returns a safe canned reply."]),
 ("K11",
  "Some agentic actions are too sensitive to fully automate — a leave approval, or a reply from a licensed "
  "financial firm to a retail client. The workflow must pause for a person to decide, and the design must also "
  "choose which steps the AI agent decides and which stay deterministic.",
  "What is human-in-the-loop in an agentic workflow, how is it implemented in n8n, and give one example of a "
  "decision that should stay deterministic rather than be given to the agent.",
  ["Human-in-the-loop: the workflow pauses until a person approves or declines before it continues.",
   "In n8n: Gmail/Outlook 'Send and Wait for Response' emails Approve/Decline buttons; an IF node routes on the decision.",
   "The agent drafts/classifies, but nothing reaches the customer until a human approves (e.g. the client-rapport assistant).",
   "Deterministic examples: trimming/normalising input, computing age from a date of birth, appending the audit-log row — "
   "never trust the agent to log itself."]),
 ("K12",
  "Voice agent platforms split the work between the vendor and your n8n workflows in different ways, and "
  "each split has its own security rules.",
  "Contrast the ElevenLabs and Vapi (Custom LLM) voice-agent architectures used in this course, and state one "
  "key security rule for each.",
  ["ElevenLabs: the vendor runs the model and the voice; n8n mints a short-lived signed URL and serves the agent's "
   "tool webhooks (e.g. check availability / book appointment).",
   "Security: the xi-api-key stays in n8n (Header Auth); the browser only ever receives the signed URL.",
   "Vapi Custom LLM: Vapi does speech-to-text and the voice, then calls YOUR n8n webhook as its model — n8n is the brain.",
   "Security: the web page holds only the Vapi PUBLIC key (it can merely start calls); the private key never goes into a page.",
   "Both vendors' SERVERS call n8n, so tool/model webhooks must be publicly reachable (n8n Cloud URL or a tunnel)."]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "FryTech Appliances is a fast-growing electronics company selling air-fryer products through e-commerce "
 "platforms, retail distributors and corporate/bulk channels. As demand grows, so does the volume and "
 "complexity of customer interactions — pre-sales enquiries, post-sales support, warranty claims and "
 "logistics. Leadership decides that simple forms and autoresponders are no longer enough and adopts "
 "agentic AI-driven workflows, orchestrated in n8n, to improve efficiency, scalability and customer "
 "satisfaction. Complete the tasks below, each mirroring an automation you built in class.")

# (label, criterion, task prompt, box caption, [model build-step points citing the activity])
PRACTICAL = [
 ("Task 1", "LO1",
  "FryTech receives a high volume of repetitive customer enquiries (pricing, wattage, capacity, warranty). "
  "Design and implement an automated customer-enquiry workflow in n8n. It must capture enquiries from a web "
  "Form, record the details in a Data Table (or Google Sheets), send a confirmation email to the customer, and "
  "notify the sales/support team — using conditional logic where appropriate. Publish the form and turn its URL "
  "into a QR code.",
  "Show the screenshot of the n8n workflow in the box below",
  ["Add a Form Trigger with fields (Name, Email, Phone, Enquiry) — as in Activity 1.",
   "Add a Gmail node to email the team and a second Gmail node to send the customer a confirmation.",
   "Add a Data Table (Insert Row) — or Google Sheets — to store each submission (Activity 2 / 3b).",
   "Optionally add an IF node to route by enquiry type / attending (Activity 3a).",
   "Save & Activate; copy the Form production URL and generate a QR code for it (Activity 1)."]),
 ("Task 2", "LO2",
  "FryTech wants a chatbot that answers team-member and customer questions about product specifications. Design "
  "and deploy an AI agent (in Telegram or embedded in a web page) that uses an LLM with memory and at least one "
  "tool to answer questions grounded in company data.",
  "Show the screenshot of the n8n workflow in the box below",
  ["Add a Telegram Trigger (Activity 4a) or a Webhook for a website chatbot (Activity 5).",
   "Add an AI Agent with a Chat Model (OpenAI/Gemini) and a Simple Memory (window buffer).",
   "Attach a tool — a Data Table tool for product/spec lookups (Activity 4b) — and write a System Instruction.",
   "Send the reply back (Telegram Send Message, or Respond to Webhook for the website).",
   "Save & Activate; test with a product question and confirm a grounded reply."]),
 ("Task 3", "LO3",
  "FryTech must integrate with external systems. Build a workflow that either (a) uses a Webhook to receive data "
  "from an external web page and respond to it, or (b) uses an HTTP Request to query an external API (e.g. market "
  "or news data) and use the result. When run, show the data is received/queried correctly.",
  "Show the screenshot of the n8n workflow in the box below",
  ["Option A — Webhook (Activity 5): add a Webhook (CORS = *), process the request, and reply with Respond to Webhook; call it from a web page.",
   "Option B — HTTP Request (Activity 6): call an external API with method, headers and query parameters; store the API key in a credential.",
   "Use the returned JSON in downstream nodes (e.g. feed it to an AI Agent).",
   "Save & Activate; run it and show the collected/queried data in the execution."]),
 ("Task 4", "LO4",
  "FryTech wants staff to query internal documents (product manuals / IT-Support FAQ / policies) in natural "
  "language. Build a Retrieval-Augmented Generation (RAG) workflow: users upload documents which are embedded "
  "and stored in a vector store, and an AI agent retrieves the most relevant chunks and answers strictly from "
  "them.",
  "Show the screenshot of the n8n workflow in the box below",
  ["Ingestion: an upload point → Embeddings → Vector Store (Insert) with a Default Data Loader (Activity 7a).",
   "Upload a document (e.g. the IT-Support FAQ PDF) from the web uploader into the vector store.",
   "Chat: an AI Agent with the Vector Store exposed as a retrieve-as-tool (knowledge_base) answers only from retrieved chunks (Activity 7a).",
   "For persistence, ingest into a vector database — Supabase / Pinecone / Qdrant — matching the embedding dimension (Activity 7b).",
   "Test: ask a question answerable only from the document; confirm a grounded answer."]),
 ("Task 5", "LO5",
  "FryTech must secure its AI agent against prompt injection and data leaks, and sensitive actions must wait "
  "for a manager's decision. Build a guarded workflow: an input guardrail that classifies each message "
  "ALLOW/BLOCK before the agent, an output guardrail on the reply, and a human-in-the-loop approval step for a "
  "sensitive action (e.g. issuing a refund or sending a customer reply).",
  "Show the screenshot of the n8n workflow in the box below",
  ["Add a Webhook → an LLM Chain 'Input Guardrail' that replies ALLOW or BLOCK (Activity 8c).",
   "Add an IF node on the result: BLOCK → Respond with a safety message; ALLOW → continue to the AI Agent.",
   "Add an 'Output Guardrail' LLM Chain (SAFE/LEAK) before Respond to Webhook (Activity 8c).",
   "Add a human-in-the-loop approval — Gmail 'Send and Wait for Response' with Approve/Decline, then an IF on "
   "the decision (Activity 8a / 8b).",
   "Test: a normal question passes; 'ignore previous instructions…' is blocked; the sensitive action waits for the manager."]),
 ("Task 6", "LO6",
  "FryTech launches a distributor onboarding programme. Applications arrive from the company website and must "
  "be screened automatically: check for an existing distributor record (duplicates), apply the eligibility "
  "rules consistently, create the record, email the applicant the outcome — and keep an audit log of every "
  "decision. Design and implement this onboarding agent, stating which steps the AI agent decides and which "
  "stay deterministic.",
  "Show the screenshot of the n8n workflow in the box below",
  ["Front door: a Webhook (POST, CORS = *) from the website — or an n8n Form for the internal version (Activity 9).",
   "Normalise the input with a Set node (trim, upper-case IDs, compute derived values) — deterministic (Activity 9).",
   "AI Agent with tools: a Sheets lookup tool for duplicates, a Sheets append tool to create the record; the "
   "eligibility rules live in the system message (Activity 9).",
   "The agent decides APPROVE / REJECT / DUPLICATE / REVIEW and writes the applicant email copy.",
   "Append the audit-log row with a plain Google Sheets node on every run — never let the agent log itself (Activity 9).",
   "Respond to Webhook returns the decision to the page; test with a clean, a duplicate and an ineligible application."]),
 ("Task 7", "LO7",
  "FryTech wants a voice assistant for its customer hotline. Build a voice agent with ElevenLabs or Vapi, "
  "connected to n8n: it must answer grounded questions about FryTech's products/policies (no invented answers) "
  "and, for the ElevenLabs option, call at least one n8n tool webhook during the conversation (e.g. check a "
  "service-appointment calendar).",
  "Show the screenshot of the n8n workflow in the box below",
  ["Option A — ElevenLabs (Activity 11): a web-call flow (Webhook → Get Signed URL with xi-api-key Header Auth → "
   "Respond) plus a tools flow (check_availability / book_appointment webhooks → Google Calendar).",
   "Upload the product/policy handbook PDF to the agent's Knowledge Base so answers are grounded.",
   "Register the tool webhooks with PUBLIC n8n URLs; the browser only ever receives the signed URL.",
   "Option B — Vapi Custom LLM (Activity 12): Webhook → AI Agent → OpenAI-shaped response; the assistant's model "
   "URL points at your n8n webhook; the page holds only the PUBLIC key.",
   "Ground the prompt in fixed FAQ topics with a fixed refusal for out-of-scope questions.",
   "Test by voice: a grounded answer, an out-of-scope refusal, and (Option A) a tool execution in n8n during the call."]),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, height_pt=90):
    """1x1 bordered box. If lines given, fill as bullet-style model answer."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO, course_code=COURSE_CODE)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO, course_code=COURSE_CODE)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    for label, crit, prompt, cap, pts in PRACTICAL:
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, lines=pts if answers else None, height_pt=150)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
