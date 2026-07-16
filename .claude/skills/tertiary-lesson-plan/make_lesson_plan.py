#!/usr/bin/env python3
"""Generate the 4-day Lesson Plan Word document for the WSQ course
'No Code and Low Code Agentic AI Applications' (TGS-2026062147)."""

import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import prodoc

REPO_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

# Logo lookup: prefer the course's .claude/skills/tertiary-course-slides/assets, else the copies bundled in this skill.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO_DIR, ".claude/skills/tertiary-course-slides/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None

# ─── EDIT PER COURSE ─────────────────────────────────────────────
TITLE       = "No Code and Low Code Agentic AI Applications"   # <<Course Title>>
COURSE_CODE = "TGS-2026062147"                                  # <<Course Code, e.g. TGS-XXXXXXXXXX>>
# ─────────────────────────────────────────────────────────────────

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LP_VERSIONS = [
    ("1.0", "17 July 2026", "First version — 4-day lesson plan for No Code and Low Code Agentic AI Applications "
                            "(TGS-2026062147). Topics 1–4 (n8n fundamentals, AI agents, webhook & HTTP request, RAG) "
                            "adapted from Agentic AI Automation with n8n; Topic 5 human-in-the-loop & guardrails; "
                            "Topic 6 use cases of agentic AI (banking onboarding, client rapport); Topic 7 voice "
                            "agents (ElevenLabs & Vapi); Topic 8 mini capstone project and presentation",
     "Tertiary Infotech Academy Pte Ltd"),
]

BRAND = RGBColor(0x1F, 0x6F, 0xEB)
DARK  = RGBColor(0x11, 0x18, 0x27)
GREY  = RGBColor(0x55, 0x5B, 0x66)
HEADER_FILL = "1F6FEB"
TOPIC_FILL  = "E8F0FE"
BREAK_FILL  = "FFF4E5"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p

def add_heading(text, size=15, color=BRAND, space_before=12, space_after=6):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.add_run(text)
    return p

# ============================ COVER + FRONT MATTER ============================
prodoc.style_headings(doc)
prodoc.add_cover_page(doc, "LESSON PLAN", TITLE, "1.0", course_code=COURSE_CODE,
    org_logo=_logo("tertiary-infotech-logo.png"),
    course_logo=_logo("n8n-course-logo.png"))
prodoc.add_version_control(doc, LP_VERSIONS)
prodoc.add_toc(doc, levels="1-2")

# ============================ COURSE INFO ============================
info = [
    ("Course Title", TITLE),
    ("Course Code", COURSE_CODE),
    ("Duration", "4 Days (32 training hours)"),
    ("Daily Schedule", "9:30 AM – 6:30 PM (8 training hours/day, excluding lunch)"),
    ("Lunch Break", "1:00 PM – 2:00 PM (1 hour)"),
    ("Breaks", "Short tea breaks are scheduled within each day's training hours"),
    ("Delivery Mode", "Instructor-led, hands-on labs with group work and presentations"),
    ("Prerequisites", "Basic computer literacy; a Google/Gmail account; no coding experience required"),
    ("Tools & Accounts", "n8n (cloud trial + local Docker), OpenAI & Gemini API keys, Telegram, "
                         "Gmail/Outlook, Google Sheets, Google Calendar, Twelve Data & NewsAPI, "
                         "ElevenLabs, Vapi"),
]
tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in info:
    row = tbl.add_row().cells
    set_cell_text(row[0], k, bold=True, color=DARK, size=10); set_cell_bg(row[0], "F1F5FB")
    set_cell_text(row[1], v, size=10)
for row in tbl.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(4.6)

# ============================ OVERVIEW ============================
add_heading("Course Overview")
doc.add_paragraph(
    "This 4-day hands-on course teaches participants to design, build and secure agentic AI "
    "applications on no-code / low-code platforms. Learners set up n8n (cloud trial and local Docker), "
    "automate forms and email, store data in Data Tables and spreadsheets, build Telegram-triggered AI "
    "agents, integrate webhooks and external APIs, add Retrieval-Augmented Generation (RAG), apply "
    "human-in-the-loop and guardrail patterns, implement two real business use cases (retail banking "
    "onboarding and a client rapport assistant with human handover), and build voice agents with "
    "ElevenLabs and Vapi — culminating in a mini capstone project that is presented and assessed."
)

# ============================ LEARNING OUTCOMES ============================
add_heading("Learning Outcomes")
doc.add_paragraph("By the end of this course, participants will be able to:")
for o in [
    "Set up n8n using a cloud trial and a local Docker Compose installation, and navigate the n8n UI.",
    "Apply the core concepts of triggers, actions, nodes and flows to build automated workflows.",
    "Automate form submissions with QR codes, send admin emails and capture data in Data Tables and Google Sheets/Excel.",
    "Design Telegram-triggered AI agents using LLMs, memory, tools and system instructions.",
    "Expose workflows via webhooks and build website chatbots that respond to live browser requests.",
    "Integrate external services through APIs and HTTP requests, including real-time market and news data.",
    "Implement Retrieval-Augmented Generation (RAG) and route an agent between multiple data sources.",
    "Apply human-in-the-loop and guardrail patterns to make AI automations safe and reliable.",
    "Apply agentic AI to business use cases — customer onboarding and client communications with human oversight.",
    "Build voice agents with ElevenLabs and Vapi connected to n8n workflows and real calendars.",
    "Plan, build, present and self-assess an end-to-end mini capstone automation project.",
]:
    p = doc.add_paragraph(o, style="List Bullet"); p.paragraph_format.space_after = Pt(2)

# ============================ SCHEDULE HELPERS ============================
def schedule_table(rows):
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Time", "Topic / Activity", "Duration"]):
        set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(hdr[i], HEADER_FILL)
    for time, activity, minutes, kind in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], time, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, bold=(kind=="topic"))
        set_cell_text(cells[1], activity, size=9.5, bold=(kind in ("topic","break")),
                      color=BRAND if kind=="topic" else (GREY if kind=="break" else None),
                      italic=(kind=="break"))
        set_cell_text(cells[2], (f"{minutes} min" if minutes else "—"), size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill = TOPIC_FILL if kind=="topic" else (BREAK_FILL if kind=="break" else None)
        if fill:
            for cc in cells: set_cell_bg(cc, fill)
    widths = [Inches(1.2), Inches(4.45), Inches(0.85)]
    for row in table.rows:
        for i, w in enumerate(widths): row.cells[i].width = w
    return table

def day_header(day, theme):
    add_heading(day, size=14, color=DARK, space_before=16)
    p = doc.add_paragraph(); r = p.add_run(theme); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(4)

def session_label(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRAND

# ----------------------------- DAY 1 -----------------------------
day_header("Day 1", "Fundamentals of n8n  ·  AI Agents")

session_label("Morning Session  ·  9:30 AM – 1:00 PM")
schedule_table([
    ("9:30 – 9:45",  f"Welcome, course objectives & overview ({COURSE_CODE})", 15, "normal"),
    ("9:45 – 12:25", "Topic 1: Fundamentals of n8n", None, "topic"),
    ("9:45 – 10:05", "Overview of n8n", 20, "normal"),
    ("10:05 – 10:35","Set up n8n — create trial accounts (15 learners) and install locally with Docker Compose", 30, "normal"),
    ("10:35 – 10:50","Explore the n8n UI", 15, "normal"),
    ("10:50 – 11:10","Key concepts of n8n — Triggers, Actions, Nodes, Flows", 20, "normal"),
    ("11:10 – 11:35","Activity 1: QR Code form (name, email, phone, message) + email to admin via Gmail/Outlook (2 nodes)", 25, "normal"),
    ("11:35 – 12:25","Group Activity (3–4 pax): design an event flyer (e.g. bowling), generate a QR code with the QR generator, build the ad, then present (review past-student examples)", 50, "normal"),
    ("12:25 – 1:00", "Activity 2: Add an n8n Data Table to Activity 1 to capture submissions (3 nodes)", 35, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])

session_label("Afternoon Session  ·  2:00 PM – 6:30 PM")
schedule_table([
    ("2:00 – 2:40",  "Activity 3a: Conditional response — if yes, save the date to the Data Table; if no, send a thank-you email (note: trial Data Table data is not persistent)", 40, "normal"),
    ("2:40 – 3:20",  "Activity 3b: Conditional response — replace the Data Table with Google Sheet/Excel for persistent storage (Google training accounts); if no, send a thank-you email", 40, "normal"),
    ("3:20 – 3:35",  "Tea Break", 15, "break"),
    ("3:35 – 6:05",  "Topic 2: AI Agents", None, "topic"),
    ("3:35 – 4:05",  "Overview of AI Agents — LLM, Memory, Tools and System Instruction", 30, "normal"),
    ("4:05 – 4:50",  "Activity 4a: Telegram-triggered AI Agent — simple customer-service chatbot (OpenAI / Gemini API key provided)", 45, "normal"),
    ("4:50 – 6:05",  "Activity 4b: Telegram-triggered AI Agent (customer service / HR admin) — attach a Data Table as a tool; reuse Activity 2 data or upload a 100-record employee CSV", 75, "normal"),
    ("6:05 – 6:30",  "Day 1 recap, Q&A and wrap-up", 25, "normal"),
])

# ----------------------------- DAY 2 -----------------------------
day_header("Day 2", "Webhook and HTTP Request  ·  RAG")

session_label("Morning Session  ·  9:30 AM – 1:00 PM")
schedule_table([
    ("9:30 – 9:45",  "Day 1 recap & Day 2 objectives", 15, "normal"),
    ("9:45 – 1:00",  "Topic 3: Webhook and HTTP Request", None, "topic"),
    ("9:45 – 10:15", "Overview of Webhooks — use cases and external triggers", 30, "normal"),
    ("10:15 – 11:35","Activity 5: Website chatbot via webhook — follow the n8n Investment Advisor reference site (CORS, Webhook trigger, Respond to Webhook, AI Agent)", 80, "normal"),
    ("11:35 – 11:50","Tea Break", 15, "break"),
    ("11:50 – 12:10","Selected learners present their live website and chatbot", 20, "normal"),
    ("12:10 – 12:40","Overview of APIs and HTTP Requests — methods, headers, query parameters, credentials", 30, "normal"),
    ("12:40 – 1:00", "Activity 6 setup: sign up for Twelve Data and NewsAPI; copy API keys; import the workflow", 20, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])

session_label("Afternoon Session  ·  2:00 PM – 6:30 PM")
schedule_table([
    ("2:00 – 3:20",  "Topic 3 (cont.): Webhook and HTTP Request", None, "topic"),
    ("2:00 – 3:20",  "Activity 6 (cont.): Configure Twelve Data and NewsAPI keys; build the Finance API → Telegram day-trading agent (candles1min / candles15min / candles1hr + news, AI Agent, Telegram reply)", 80, "normal"),
    ("3:20 – 6:30",  "Topic 4: Retrieval-Augmented Generation (RAG)", None, "topic"),
    ("3:20 – 4:00",  "Overview of Retrieval-Augmented Generation (RAG) — Tokenization, Embeddings, Vector Stores", 40, "normal"),
    ("4:00 – 4:15",  "Tea Break", 15, "break"),
    ("4:15 – 4:45",  "RAG ingestion walkthrough — the upload → embed → vector-store → retrieve loop; set up OpenAI / Gemini embedding credentials", 30, "normal"),
    ("4:45 – 6:05",  "Activity 7a: RAG chatbot — upload an IT-Support FAQ PDF from a web page into a vector store, then ask the Telegram bot (Gemini embeddings + knowledge_base tool).  Activity 7b: customer-support RAG agent for a training center, ingesting 20 course brochures across three vector databases (Supabase pgvector, Pinecone, Qdrant) answering via a website chat widget", 80, "normal"),
    ("6:05 – 6:30",  "Day 2 recap, Q&A and wrap-up", 25, "normal"),
])

# ----------------------------- DAY 3 -----------------------------
day_header("Day 3", "Human in the Loop and Guardrails  ·  Use Cases of Agentic AI")

session_label("Morning Session  ·  9:30 AM – 1:00 PM")
schedule_table([
    ("9:30 – 9:45",  "Day 2 recap & Day 3 objectives", 15, "normal"),
    ("9:45 – 1:00",  "Topic 5: Human in the Loop and Guardrails — HR Service Portal", None, "topic"),
    ("9:45 – 10:25", "Overview of Human-in-the-Loop and the HR Service Portal — three coordinated n8n workflows behind one web page (Leave Application, Dashboard, AI Chatbot)", 40, "normal"),
    ("10:25 – 11:25","Activity 8a: Leave Application with Human-in-the-Loop — webhook receives leave request, emails manager an Approve/Reject button (Send and Wait for Response), emails employee the outcome", 60, "normal"),
    ("11:25 – 11:40","Tea Break", 15, "break"),
    ("11:40 – 12:00","Selected learners present their Activity 8a approval flow", 20, "normal"),
    ("12:00 – 12:25","Overview of Guardrails — pre/post patterns, blocked responses, canned-reply fallbacks", 25, "normal"),
    ("12:25 – 1:00", "Activity 8b & 8c: Build the HR Dashboard Data webhook (leave-balance stats) and the AI Chatbot with Input & Output Guardrails; wire all three endpoints in the HR Service Portal", 35, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])

session_label("Afternoon Session  ·  2:00 PM – 6:30 PM")
schedule_table([
    ("2:00 – 6:30",  "Topic 6: Use Cases of Agentic AI", None, "topic"),
    ("2:00 – 2:20",  "Overview — agentic AI in the workplace; the 4-step problem-solving framework; agent decisions vs deterministic steps", 20, "normal"),
    ("2:20 – 2:55",  "Activity 9 Part A (group): retail banking onboarding — problem definition, root cause analysis, solution ideation (Marina Trust Bank)", 35, "normal"),
    ("2:55 – 4:15",  "Activity 9 Part B: build the onboarding AI Agent (duplicate check, KYC screening, eligibility rules, customer record, email) — form version, then the bank-website webhook version", 80, "normal"),
    ("4:15 – 4:30",  "Tea Break", 15, "break"),
    ("4:30 – 5:45",  "Activity 10: client rapport assistant with human handover (Meridian Asset Management) — agent drafts a non-advisory reply; a relationship manager approves or declines before anything is sent; three-tab audit trail", 75, "normal"),
    ("5:45 – 6:05",  "Selected learners present their use-case builds", 20, "normal"),
    ("6:05 – 6:30",  "Day 3 recap, Q&A and wrap-up", 25, "normal"),
])

# ----------------------------- DAY 4 -----------------------------
day_header("Day 4", "Voice Agents  ·  Mini Capstone Project and Presentation")

session_label("Morning Session  ·  9:30 AM – 1:00 PM")
schedule_table([
    ("9:30 – 9:45",  "Day 3 recap & Day 4 objectives", 15, "normal"),
    ("9:45 – 1:00",  "Topic 7: Voice Agents", None, "topic"),
    ("9:45 – 10:15", "Overview of Voice Agents — ElevenLabs vs Vapi architectures; signed URLs; public vs private keys; tool webhooks", 30, "normal"),
    ("10:15 – 12:15","Activity 11: Voice booking agent with ElevenLabs (GG Hair Salon) — signed-URL flow, Google Calendar tool webhooks, knowledge-base grounding, end-to-end voice booking", 120, "normal"),
    ("12:15 – 12:30","Tea Break", 15, "break"),
    ("12:30 – 1:00", "Activity 12 setup: import the Vapi Custom-LLM flow, prove the webhook with curl, read the guardrail prompt", 30, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])

session_label("Afternoon Session  ·  2:00 PM – 6:30 PM")
schedule_table([
    ("2:00 – 3:00",  "Topic 7 (cont.): Voice Agents", None, "topic"),
    ("2:00 – 3:00",  "Activity 12 (cont.): grounded FAQ voice agent with Vapi (MediRefill) — create the assistant, wire the public key, run the graded calls (grounded answer, medical refusal, emergency escalation)", 60, "normal"),
    ("3:00 – 6:30",  "Topic 8: Mini Capstone Project and Presentation", None, "topic"),
    ("3:00 – 3:20",  "Capstone briefing — requirements, scope and assessment criteria", 20, "normal"),
    ("3:20 – 4:35",  "Mini capstone build (supervised lab time)", 75, "normal"),
    ("4:35 – 4:50",  "Tea Break", 15, "break"),
    ("4:50 – 5:45",  "Presentation of mini capstone projects", 55, "normal"),
    ("5:45 – 6:30",  "Assessment (WA SAQ + PP briefing) and course closing", 45, "normal"),
])

# ============================ RESOURCES ============================
add_heading("Tools, Accounts & Resources", space_before=16)
res = [
    ("QR Code Generator", "https://alfredang.github.io/qrcodegenerator/"),
    ("Webhook reference — Activity 5: Investment Advisor", "https://alfredang.github.io/n8n-investmentadvisor/"),
    ("Finance API reference — Activity 6: Finance Advisor", "https://alfredang.github.io/n8n-financeadvisor/"),
    ("Finance Data API — Twelve Data", "https://twelvedata.com/login"),
    ("News API — NewsAPI", "https://newsapi.org/"),
    ("LLM API keys", "OpenAI API key & Google Gemini API key (provided to learners)"),
    ("Voice — ElevenLabs", "https://elevenlabs.io/ (Conversational AI agent + API key)"),
    ("Voice — Vapi", "https://vapi.ai/ (assistant + public key)"),
    ("Local n8n", "Docker Compose self-hosted install (see Learner Guide)"),
]
rt = doc.add_table(rows=0, cols=2); rt.style = "Table Grid"; rt.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in res:
    row = rt.add_row().cells
    set_cell_text(row[0], k, bold=True, color=DARK, size=9.5); set_cell_bg(row[0], "F1F5FB")
    set_cell_text(row[1], v, size=9.5, color=BRAND)
for row in rt.rows:
    row.cells[0].width = Inches(2.4); row.cells[1].width = Inches(4.1)

# ============================ ASSESSMENT ============================
add_heading("Assessment", space_before=14)
doc.add_paragraph(
    "Participants are assessed through a Written Assessment (SAQ, 1 hour) and a Practical "
    "Performance assessment (1 hour), plus hands-on activities across the four days and a mini "
    "capstone project presented on Day 4. Work is evaluated against the course learning outcomes — "
    "workflow design, AI agent integration, webhook and API use, Retrieval-Augmented Generation, "
    "human-in-the-loop and guardrails, business use cases, and voice agents."
)

# ============================ FOOTER (page numbers) ============================
prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)

OUT = os.path.join(REPO_DIR, f"courseware/LP-{TITLE}.docx")
doc.save(OUT)
print("Saved:", OUT)

# ---- sanity check: every day = 480 min (incl. tea breaks, excl. lunch) ----
days = {
    "Day 1": [15,20,30,15,20,25,50,35] + [40,40,15,30,45,75,25],
    "Day 2": [15,30,80,15,20,30,20] + [80,40,15,30,80,25],
    "Day 3": [15,40,60,15,20,25,35] + [20,35,80,15,75,20,25],
    "Day 4": [15,30,120,15,30] + [60,20,75,15,55,45],
}
for d, mins in days.items():
    total = sum(mins)
    print(f"{d}: {total} min = {total/60:.2f} hrs ->", "OK" if total == 480 else f"MISMATCH (expected 480)")
