#!/usr/bin/env python3
"""
Single-source generator for the n8n Learner Guide.
Emits BOTH:
  - LEARNER-GUIDE.md                                  (repo root)
  - courseware/n8n Automation Learner Guide.docx
from one content model so the two are always aligned.

Content DSL (list of tuples):
  ("h1", text) ("h2", text) ("h3", text)
  ("p", text)
  ("steps", [..])      numbered list
  ("bullets", [..])    bullet list
  ("code", text)       fenced/monospace block
  ("table", [hdr,...rows])
  ("note", text)       callout
  ("rule",)            horizontal divider
"""
import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = "/Users/alfredang/projects/labs-activities/TGS-2023035977-Agentic AI Automation with n8n"

# ============================================================================
# CONTENT
# ============================================================================
B = []
def h1(t): B.append(("h1", t))
def h2(t): B.append(("h2", t))
def h3(t): B.append(("h3", t))
def p(t):  B.append(("p", t))
def steps(xs): B.append(("steps", xs))
def bullets(xs): B.append(("bullets", xs))
def code(t): B.append(("code", t))
def table(rows): B.append(("table", rows))
def note(t): B.append(("note", t))
def rule(): B.append(("rule",))

# ---------------------------------------------------------------- Title / intro
h1("Agentic AI Automation with n8n — Step-by-Step Learner Guide")
p("Welcome! This guide takes you click-by-click through every hands-on lab in the WSQ course "
  "**Agentic AI Automation with n8n** (Course Code: TGS-2023035977). Over three days you will go from "
  "simple form automations to AI agents, Retrieval-Augmented Generation (RAG), webhooks, external APIs, "
  "and finally human-in-the-loop guardrails — then build a mini capstone of your own.")
p("Work through the activities in order: each one builds on the skills (and sometimes the workflow) of the "
  "activity before it. Whenever you see a **Test it** box, stop and confirm your workflow behaves as described "
  "before moving on.")

note("Course flow at a glance — "
     "Day 1: Workflow Automation (Activities 1-3) + AI Agents (Activity 4). "
     "Day 2: RAG (Activity 5) · Webhooks (Activity 6) · APIs (Activity 7). "
     "Day 3: Security & Guardrails (Activity 8) + Mini Capstone.")

# ---------------------------------------------------------------- 0. Setup
rule()
h2("0. Before You Start — Setup & Prerequisites")

h3("0.1 Accounts & API keys you will need")
table([
    ["Service", "Used for", "Where to get it"],
    ["n8n", "The automation platform (all activities)", "Cloud trial at n8n.io, or local Docker (see 0.2)"],
    ["Gmail or Outlook", "Sending emails (Activities 1-3)", "Your existing mailbox; connected via OAuth2"],
    ["OpenAI API key", "LLM for AI agents (Activities 4-8)", "platform.openai.com/api-keys (provided in class)"],
    ["Google Gemini API key", "Alternative LLM", "aistudio.google.com/app/apikey (provided in class)"],
    ["Telegram", "Chat trigger for AI agents (Activities 4,5,7)", "Telegram app + @BotFather"],
    ["Google account", "Google Sheets storage (Activity 3b)", "Your Google account (training account provided)"],
    ["Twelve Data", "Live market data (Activity 7)", "twelvedata.com — free API key"],
    ["NewsAPI", "Headlines & sentiment (Activity 7)", "newsapi.org — free API key"],
])

h3("0.2 Run n8n — Cloud trial OR local Docker")
p("You have two ways to run n8n. For the course we start everyone on a **cloud trial** so we are productive "
  "immediately, and we also show how to **self-host locally with Docker** so you can keep your workflows after "
  "the trial ends.")
p("**Option A — Cloud trial (fastest).** Sign up at n8n.io, create a workspace, and you land directly in the "
  "workflow editor. Note: trial **Data Tables are not permanent** — for anything you want to keep, store it "
  "externally (e.g. Google Sheets, see Activity 3b).")
p("**Option B — Local install with Docker Compose (persistent).** Install Docker Desktop, then create a file "
  "named `docker-compose.yml` (a ready-made copy is in `labs/n8n-installation/`):")
code(
"# labs/n8n-installation/docker-compose.yml\n"
"services:\n"
"  n8n:\n"
"    image: docker.n8n.io/n8nio/n8n\n"
"    restart: always\n"
"    ports:\n"
"      - \"5678:5678\"\n"
"    environment:\n"
"      - N8N_SECURE_COOKIE=false\n"
"      - GENERIC_TIMEZONE=Asia/Singapore\n"
"    volumes:\n"
"      - n8n_data:/home/node/.n8n\n"
"volumes:\n"
"  n8n_data:")
steps([
    "Open a terminal in the `labs/n8n-installation/` folder.",
    "Run `docker compose up -d` to start n8n in the background.",
    "Open http://localhost:5678 in your browser and create your owner account.",
    "Your workflows and credentials now persist in the `n8n_data` volume, even after restarts.",
    "To stop n8n, run `docker compose down` (your data is kept); to update, `docker compose pull` then `up -d` again.",
])

h3("0.3 Add your credentials in n8n (do this once)")
p("Credentials are stored separately from workflows so you never paste secrets into nodes. Add them under "
  "**Credentials → Add credential**:")
bullets([
    "**Gmail / Microsoft Outlook (OAuth2)** — sign in and authorise n8n to send mail on your behalf.",
    "**OpenAI** — paste your OpenAI API key. (Gemini: add a *Google Gemini (PaLM) API* credential instead.)",
    "**Telegram** — paste the bot token from @BotFather (see Activity 4a, Step 1).",
    "**Google Sheets (OAuth2)** — authorise access to your Google Sheets (Activity 3b).",
    "**HTTP Header Auth / query params** — for Twelve Data & NewsAPI keys (Activity 7).",
])
note("Imported workflows reference credential *names*, not your actual secrets. After importing any provided "
     "`.json`, re-select your own credentials on each node that needs them.")

h3("0.4 Download the workflows from GitHub")
p("All the finished workflow `.json` files, the mock data (CSV) and the sample documents are in the course "
  "GitHub repository — download them so you can import and follow along:")
note("**GitHub repo:** https://github.com/tertiarycourses/TGS-2023035977-Agentic-AI-Automation-with-n8n  ·  "
     "every activity lives under the **`labs/`** folder (one folder per activity), each with its workflow JSON, "
     "a workflow diagram, and any mock data.")
steps([
    "Open the repo and click **Code → Download ZIP** (or `git clone` it).",
    "Each activity folder under `labs/` contains the importable workflow `.json` and its mock data.",
    "In n8n, open the **Workflows** list → **Add workflow** → the **⋯** menu → **Import from File**.",
    "Choose the matching `.json`, then re-select your own credentials on each node (OpenAI, Gmail, Telegram, etc.).",
    "**Save**, then toggle the workflow **Active** when the activity says to.",
])

# ============================================================================
# DAY 1
# ============================================================================
rule()
h2("Activity 1 — Flyer with QR Code (Form → Email)")
p("**Folder:** `labs/activity1-flyer-form/`")
h3("Goal")
p("Build the smallest useful automation: an n8n **Form** that collects a visitor's details and emails them to "
  "an admin. You will then turn the form's URL into a **QR code** and put it on an event flyer.")
h3("What you'll build (2 nodes)")
p("**Form Trigger** → **Gmail (Send)**")
h3("Step-by-step")
steps([
    "Create a new workflow and name it `Activity 1 — Flyer Form`.",
    "Add an **n8n Form Trigger** node. Set a **Form Title** (e.g. \"Event RSVP\").",
    "Add four form fields: **Name** (text), **Email** (email), **Phone** (text), **Message** (textarea). Mark Name and Email required.",
    "Add a **Gmail** node, operation **Send a Message**, connected after the Form Trigger.",
    "In the Gmail node set **To** = the admin address, **Subject** = `New RSVP from {{ $json.Name }}`.",
    "Set the message body to include the submitted fields, e.g. `Name: {{ $json.Name }} / Email: {{ $json.Email }} / Phone: {{ $json.Phone }} / Message: {{ $json.Message }}`.",
    "Select your Gmail (or Outlook) credential. **Save** the workflow and toggle it **Active**.",
    "Open the Form Trigger node and copy the **Production URL**.",
])
h3("Make the QR code & flyer")
steps([
    "Open the QR code generator: https://alfredang.github.io/qrcodegenerator/",
    "Paste your form's Production URL and generate the QR code; download it.",
    "Place the QR code on a flyer so people can scan to open your form.",
])
note("Group Activity (3-4 per group): design a real event flyer — e.g. a bowling night — with your form's QR "
     "code and a short advert. Review a few past-student examples first, then present your flyer to the class.")
B.append(("test", "Scan the QR code with your phone, submit the form, and confirm the admin inbox receives the email."))

rule()
h2("Activity 2 — Capture Submissions in a Data Table")
p("**Folder:** `labs/activity2-data-table/`")
h3("Goal")
p("Extend Activity 1 so every submission is also **saved** into an n8n **Data Table** — your first taste of "
  "storing data, not just forwarding it.")
h3("What you'll build (3 nodes)")
p("**Form Trigger** → **Gmail (Send)** and **Data Table (Insert row)**")
h3("Step-by-step")
steps([
    "In n8n open **Data Tables** and create a table named `RSVPs` with columns: Name, Email, Phone, Message.",
    "Duplicate your Activity 1 workflow (or continue in it).",
    "Add a **Data Table** node, operation **Insert Row**, and connect it after the Form Trigger (alongside Gmail).",
    "Map each form field to the matching column using expressions, e.g. Name → `{{ $json.Name }}`.",
    "**Save** and keep the workflow **Active**.",
])
B.append(("test", "Submit the form again and confirm a new row appears in the `RSVPs` Data Table and the email still sends."))

rule()
h2("Activity 3a — Conditional Response (Data Table)")
p("**Folder:** `labs/activity3-conditional/`")
h3("Goal")
p("Add decision-making. Ask \"Will you attend?\" — if **Yes**, save the date to the Data Table; if **No**, send "
  "a polite thank-you email instead.")
h3("What you'll build")
p("**Form Trigger** → **IF** → (true) **Data Table Insert** / (false) **Gmail thank-you**")
h3("Step-by-step")
steps([
    "Add an **Attending?** field to the form (dropdown: Yes / No).",
    "Add an **IF** node after the Form Trigger. Condition: `{{ $json.Attending }}` **equals** `Yes`.",
    "On the **true** output, add a **Data Table → Insert Row** node that saves the RSVP (Name, Email, Date).",
    "On the **false** output, add a **Gmail** node that sends a friendly \"thanks anyway\" message.",
    "**Save** and keep the workflow **Active**.",
])
note("Heads-up on persistence: data saved to a trial **Data Table disappears when the trial ends**. To keep it "
     "permanently you must store it **externally** — that is exactly what Activity 3b does with Google Sheets.")
B.append(("test", "Submit once with Attending = Yes (expect a new Data Table row) and once with No (expect the thank-you email)."))

rule()
h2("Activity 3b — Conditional Response (Google Sheets / Excel)")
p("**Folder:** `labs/activity3-conditional/`")
h3("Goal")
p("Make your data **persistent** by replacing the Data Table with **Google Sheets** (or Excel). Same logic as "
  "3a, but the \"Yes\" branch now appends a row to a real spreadsheet you keep.")
h3("What you'll build")
p("**Form Trigger** → **IF** → (true) **Google Sheets (Append row)** / (false) **Gmail thank-you**")
h3("Step-by-step")
steps([
    "In Google Drive, create a spreadsheet named `Event RSVPs` with a header row: Name, Email, Phone, Date.",
    "Back in n8n, add a **Google Sheets** credential (OAuth2) and authorise it.",
    "Take your Activity 3a workflow and on the **true** branch replace the Data Table node with a "
    "**Google Sheets → Append Row** node.",
    "Select your spreadsheet and sheet, then map each column to the form fields.",
    "Leave the **false** branch (thank-you email) unchanged. **Save** and keep **Active**.",
])
note("Microsoft 365 users can use the **Microsoft Excel 365** node instead of Google Sheets — the steps are the same.")
B.append(("test", "Submit with Attending = Yes and confirm a new row is appended to your Google Sheet."))

rule()
h2("Activity 4a — Telegram-Triggered AI Agent (Customer Service)")
p("**Folder:** `labs/activity4-telegram-agent/`")
h3("Goal")
p("Build your first **AI Agent**: a simple customer-service chatbot you talk to from **Telegram**. The agent "
  "uses an LLM, short-term memory, and a system instruction that defines its persona.")
h3("Concepts — what makes an AI Agent")
bullets([
    "**LLM** — the model that generates replies (OpenAI `gpt-4.1-mini` or Google Gemini).",
    "**Memory** — remembers the recent conversation so follow-up questions make sense.",
    "**Tools** — optional actions the agent can call (added in Activity 4b).",
    "**System Instruction** — the agent's role, tone, and rules.",
])
h3("Step 1 — Create the Telegram bot")
steps([
    "In Telegram, open **@BotFather** → `/newbot`, give it a name and username, and copy the **bot token**.",
    "In n8n add a **Telegram** credential and paste the token.",
])
h3("Step 2 — Build the workflow")
steps([
    "Add a **Telegram Trigger** node (it fires on each incoming message). Select your Telegram credential.",
    "Add an **AI Agent** node connected after the trigger.",
    "Attach an **OpenAI Chat Model** (or **Google Gemini Chat Model**) as the agent's model.",
    "Attach a **Simple Memory** node so the agent recalls the conversation.",
    "Write the **System Instruction**, e.g. \"You are a friendly customer-service assistant for MyCompany. "
    "Answer concisely and politely.\"",
    "Add a **Telegram → Send Message** node after the agent; set **Chat ID** = `{{ $json.message.chat.id }}` and "
    "**Text** = the agent's output.",
    "**Save** and toggle **Active**.",
])
B.append(("test", "Message your bot in Telegram (e.g. \"What are your opening hours?\") and confirm it replies."))

rule()
h2("Activity 4b — Telegram Agent + Data Table Tool (HR Admin)")
p("**Folder:** `labs/activity4-telegram-agent/`")
h3("Goal")
p("Give the agent a **tool**: an employee **Data Table** it can look up. Now the same Telegram bot can answer "
  "HR-admin questions like \"What department is Alice in?\" by querying real data.")
h3("Step-by-step")
steps([
    "Create a Data Table named `Employees`. Either reuse data from Activity 2, or upload the provided "
    "`mock-hr-employees.csv` (100 records) — regenerate it any time with `make_mock_data.py`.",
    "Open your Activity 4a workflow.",
    "Add a **Data Table Tool** and attach it to the AI Agent's **Tool** input.",
    "Point the tool at the `Employees` table and describe it in the tool description, e.g. "
    "\"Look up employee details by name or department.\"",
    "Update the **System Instruction**: \"You are an HR admin assistant. Use the Employees tool to answer "
    "questions about staff. If the data is not found, say so.\"",
    "**Save** and keep **Active**.",
])
B.append(("test", "Ask the bot \"Which department is <a name from the CSV> in?\" and confirm it answers from the table."))

rule()
h2("Activity 5 — Add RAG to the Telegram Agent (Two Knowledge Sources)")
p("**Folder:** `labs/activity5-rag/`")
h3("Goal")
p("Upgrade the agent with **Retrieval-Augmented Generation (RAG)** so it can answer from **documents** (policy "
  "PDFs/FAQs) as well as the **Data Table**. The agent must route to the **right source** for each question.")
h3("Concepts — RAG in one minute")
bullets([
    "**Tokenization** — text is split into tokens the model can process.",
    "**Embeddings** — each chunk of a document becomes a vector (a list of numbers capturing meaning).",
    "**Vector store** — those vectors are saved so the most relevant chunks can be retrieved for a question.",
])
h3("Step-by-step")
steps([
    "Prepare knowledge documents. Use the provided `MyCompany-HR-SOP.docx` and `MyCompany-IT-Support-FAQ.docx`, "
    "or generate fresh ones with Claude Code (e.g. an employee-benefits FAQ or product info).",
    "Build the **ingestion** path: an upload point → **Embeddings (OpenAI)** → **Vector Store (Insert)** with a "
    "**Default Data Loader** so your documents are embedded and stored.",
    "In your Telegram agent, add a **Vector Store** retrieval **tool** (the RAG source) alongside the existing "
    "**Data Table** tool — the agent now has **two** data sources.",
    "Rewrite the **System Instruction** to route correctly: \"Use the **Knowledge Base** tool for policy/FAQ "
    "questions and the **Employees** tool for staff-record questions. Never mix the two.\"",
    "**Save** and keep **Active**.",
])
note("Get a few learners to present their chatbot and show it answering both a policy question (RAG) and a "
     "staff-record question (Data Table).")
B.append(("test", "Ask a policy question (\"How many days of annual leave do I get?\") and a record question "
                  "(\"What is Alice's role?\") and confirm each is answered from the correct source."))

# ============================================================================
# DAY 2 (Webhook + API)
# ============================================================================
rule()
h2("Activity 6 — Website Chatbot via Webhook (Investment Advisor)")
p("**Folder:** `labs/activity6-investment-advisor/`  ·  Reference: https://alfredang.github.io/n8n-investmentadvisor/")
h3("Goal")
p("Expose an AI agent to a **public website** using a **Webhook**. The provided one-page Investment Advisor "
  "site has an enquiry form and a floating chatbot; both POST to a single n8n webhook, which routes to an "
  "email-the-advisor path and an AI-chat path.")
h3("Concepts — Webhooks")
bullets([
    "A **Webhook** is a URL that external systems (a website, another app) call to **trigger** your workflow.",
    "Use cases: website chat, form submissions, payment events, GitHub/Stripe notifications — any external trigger.",
    "Pair the Webhook trigger with a **Respond to Webhook** node to send a reply back to the caller.",
])
h3("Step-by-step")
steps([
    "Import `Activity6-Investment-Advisor.json` into n8n.",
    "Open the **Webhook** node(s) and ensure **Allowed Origins (CORS)** is `*` so the browser page can call it.",
    "Re-select your **OpenAI** and **Gmail** credentials on the AI Agent and Email nodes.",
    "Review the agent's compliance system instruction (no guaranteed returns, no personalised advice).",
    "**Save**, toggle **Active**, and copy the webhook **Production URL**.",
    "Open `index.html` from the activity folder, click the settings/gear, and paste your webhook URL.",
])
note("Get a few learners to present their live website and chatbot.")
B.append(("test", "On the website, send a chat message and submit the enquiry form; confirm the bot replies and the advisor receives the enquiry email."))

rule()
h2("Activity 7 — Finance API → Telegram (AI Day-Trading Agent)")
p("**Folder:** `labs/activity7-finance-advisor/`  ·  Reference: https://alfredang.github.io/n8n-financeadvisor/")
h3("Goal")
p("Combine **APIs/HTTP Requests** with an AI agent. Ask the Telegram bot about a stock; it resolves the ticker, "
  "pulls **multi-timeframe candles from Twelve Data** and **headlines from NewsAPI**, and replies with a "
  "Buy / Sell / Hold call and reasoning. A companion dashboard shows live price and charts.")
h3("Concepts — APIs & HTTP Request")
bullets([
    "An **API** lets your workflow request data from another service over HTTP.",
    "The **HTTP Request** node calls an endpoint with a method (GET/POST), headers, and query parameters.",
    "**API keys** authenticate you — keep them in credentials, never hard-coded.",
])
h3("Get your API keys")
steps([
    "Twelve Data — sign in at https://twelvedata.com/login and copy your free API key.",
    "NewsAPI — register at https://newsapi.org/ and copy your API key.",
])
h3("Step-by-step")
steps([
    "Import `Activity7-Finance-Advisor.json` into n8n.",
    "Set your **Twelve Data** and **NewsAPI** keys (in credentials or the HTTP Request query params).",
    "Re-select your **OpenAI** and **Telegram** credentials.",
    "Review the flow: **Telegram Trigger → Extract Ticker (LLM) → HTTP candles (1m/15m/1h) + HTTP news → "
    "Aggregate/Merge → AI Agent → Telegram reply**.",
    "**Save**, toggle **Active**. Optionally open `index.html`, set your Twelve Data key and bot username in "
    "the dashboard settings.",
])
B.append(("test", "Message the bot \"Should I buy AAPL?\" and confirm it returns a recommendation with reasoning."))

# ============================================================================
# DAY 3 (Security + Capstone)
# ============================================================================
rule()
h2("Activity 8a — Human-in-the-Loop Approval (Leave Application)")
p("**Folder:** `labs/activity8-guardrails/`")
h3("Goal")
p("Add a **human approval** step so the automation pauses for a person to decide. We model a **leave-application "
  "approval**: a request comes in, a manager is asked to approve, and the flow only continues on approval.")
h3("Concepts — Human in the Loop")
bullets([
    "Some actions are too sensitive to fully automate — money, hiring, sending on someone's behalf.",
    "A **human-in-the-loop** step pauses the workflow and waits for a person to **Approve** or **Reject**.",
    "n8n provides **Send and Wait for Response** (e.g. via email/Telegram) to capture that decision.",
])
h3("Step-by-step")
steps([
    "Start a workflow with a **Form Trigger** (or Telegram) collecting: Employee, Dates, Reason.",
    "Add a **Gmail → Send and Wait for Response** (Approval) node addressed to the manager, with **Approve** / "
    "**Reject** buttons.",
    "On **Approved**, record the leave (Data Table or Google Sheet) and email a confirmation to the employee.",
    "On **Rejected**, email the employee that the request was declined.",
    "**Save** and keep **Active**.",
])
note("Get a few learners to present their approval flow.")
B.append(("test", "Submit a leave request, approve it from the manager email, and confirm the employee gets a confirmation."))

rule()
h2("Activity 8b — Pre & Post Guardrails for the AI Agent")
p("**Folder:** `labs/activity8-guardrails/`")
h3("Goal")
p("Wrap an AI agent with **guardrails** so unsafe input never reaches the model and unsafe output never reaches "
  "the user. You add a **pre-check** before the agent and a **post-check** after it.")
h3("Concepts — Guardrails")
bullets([
    "**Pre-guardrail** — validate/sanitise the *input* (block prompt-injection, PII, banned topics) before the LLM.",
    "**Post-guardrail** — check the *output* (no secrets, no disallowed content) before it is sent.",
    "If a guardrail fails, route to a safe fallback (a canned reply, or human review).",
])
h3("Step-by-step")
steps([
    "Take the Activity 6 webhook agent (or the Telegram agent).",
    "**Before** the AI Agent, add a check node (a Guardrails node, or an LLM/If classifier) that inspects the "
    "user message; on a violation, branch to a safe canned response instead of the agent.",
    "**After** the AI Agent, add a second check that scans the reply for secrets/policy violations; on a "
    "violation, replace it with a safe message (or send it for human review as in 8a).",
    "Only send the reply to the user when both guardrails pass.",
    "**Save** and keep **Active**.",
])
B.append(("test", "Send a normal question (passes through) and a disallowed one (blocked by the pre-guardrail with a safe reply)."))

rule()
h2("Mini Capstone Project")
p("**Folder:** `labs/mini-capstone/`")
h3("Goal")
p("Bring it together. In small groups, design and build an end-to-end automation that uses what you learned: a "
  "trigger (form/Telegram/webhook), an AI agent with at least one **tool** or **RAG** source, an external "
  "**API** or storage, and a **guardrail** or human-in-the-loop step. A worked example — an **Issue Reporting** "
  "flow (form + image → Postgres, with a retrieval API and gallery) — is provided in the folder.")
h3("Deliverables")
bullets([
    "A working n8n workflow (exported `.json`).",
    "A short demo of the happy path and at least one safety/guardrail case.",
    "A 3-5 minute presentation: problem, design, what you'd improve.",
])
h3("Assessment")
p("Your capstone and the activities across the three days are assessed against the course learning outcomes: "
  "workflow design, AI agent / RAG integration, webhook & API use, and the application of security guardrails.")

# ---------------------------------------------------------------- Troubleshooting
rule()
h2("Troubleshooting Cheat-Sheet")
table([
    ["Symptom", "Likely cause & fix"],
    ["Browser page can't reach the webhook", "Set the Webhook node's **Allowed Origins (CORS)** to `*`; use the **Production** URL with the workflow Active."],
    ["Imported workflow errors on run", "Re-select your own credentials on every node; imported credential IDs won't match."],
    ["AI agent gives empty/odd replies", "Check the model credential is valid and the System Instruction is set; confirm Memory is attached."],
    ["Telegram bot doesn't respond", "Workflow must be **Active**; the Telegram credential token must match the bot; check the chat ID expression."],
    ["Data Table data disappeared", "Trial Data Tables are not permanent — use Google Sheets/Excel (Activity 3b) for persistence."],
    ["API returns 401/429", "401 = wrong/missing API key; 429 = rate limit — wait, or reduce request frequency."],
])

# ---------------------------------------------------------------- Glossary
rule()
h2("Glossary")
table([
    ["Term", "Meaning"],
    ["Trigger", "The node that starts a workflow (Form, Webhook, Telegram, Schedule, Manual)."],
    ["Node", "A single step/block in a workflow (an action, a logic gate, a trigger)."],
    ["Action", "A node that does something — send email, insert a row, call an API."],
    ["Flow / Connection", "The wires linking nodes, defining execution order and data passing."],
    ["AI Agent", "A node that uses an LLM plus memory and tools to reason and act."],
    ["LLM", "Large Language Model — the AI that understands and generates text."],
    ["RAG", "Retrieval-Augmented Generation — answering from your documents via a vector store."],
    ["Embedding", "A numeric vector representing the meaning of a piece of text."],
    ["Vector store", "A database of embeddings used to retrieve relevant chunks."],
    ["Webhook", "A URL that external systems call to trigger a workflow."],
    ["Guardrail", "A safety check on an agent's input (pre) or output (post)."],
    ["Human in the loop", "A pause for a person to approve/reject before the flow continues."],
])

p("You're done — congratulations! Keep your local n8n running to continue building your own agents.")

# ============================================================================
# RENDERERS
# ============================================================================
TITLE = "Agentic AI Automation with n8n"
VERSION = "3.0"
VERSIONS = [
    ("1.0", "2 Feb 2023", "First version", "Dr. Alfred Ang"),
    ("2.0", "16 June 2025", "Updated course title and content", "Tertiary Infotech Pte Ltd"),
    ("3.0", "24 June 2026", "Restructured to 8 activities; aligned to the agentic n8n flow "
                            "(Telegram agents, RAG, webhooks, APIs, guardrails); MD and DOCX aligned",
     "Tertiary Infotech Academy Pte Ltd"),
]

def _toc(blocks):
    lines = ["## Table of Contents", ""]
    for b in blocks:
        if b[0] == "h2":
            anchor = b[1].lower().replace("—", "").replace("(", "").replace(")", "")
            anchor = "-".join(anchor.split()).replace("/", "").replace(".", "").replace(",", "")
            lines.append(f"- [{b[1]}](#{anchor})")
    lines.append("")
    return "\n".join(lines)

def render_markdown(blocks):
    out = []
    # version control + TOC injected right after the first H1
    injected = False
    for b in blocks:
        k = b[0]
        if k == "h1":
            out.append(f"# {b[1]}\n")
            if not injected:
                out.append("**Course Code:** TGS-2023035977  ·  **Version 3.0**  ·  Tertiary Infotech Academy Pte Ltd\n")
                out.append("### Document Version Control Record\n")
                out.append("| Version | Effective Date | Summary of Changes | Author |")
                out.append("| --- | --- | --- | --- |")
                for v in VERSIONS:
                    out.append(f"| {v[0]} | {v[1]} | {v[2]} | {v[3]} |")
                out.append("")
                out.append(_toc(blocks))
                injected = True
            continue
        if k == "h1": out.append(f"# {b[1]}\n")
        elif k == "h2": out.append(f"## {b[1]}\n")
        elif k == "h3": out.append(f"### {b[1]}\n")
        elif k == "p": out.append(f"{b[1]}\n")
        elif k == "steps":
            out.append("\n".join(f"{i}. {s}" for i, s in enumerate(b[1], 1)) + "\n")
        elif k == "bullets":
            out.append("\n".join(f"- {s}" for s in b[1]) + "\n")
        elif k == "code":
            out.append("```\n" + b[1] + "\n```\n")
        elif k == "table":
            rows = b[1]; hdr = rows[0]
            out.append("| " + " | ".join(hdr) + " |")
            out.append("| " + " | ".join("---" for _ in hdr) + " |")
            for r in rows[1:]:
                out.append("| " + " | ".join(r) + " |")
            out.append("")
        elif k == "img":
            out.append(f"![{b[2]}]({b[1]})\n")
            out.append(f"*{b[2]}*\n")
        elif k == "note":
            out.append(f"> **Note:** {b[1]}\n")
        elif k == "test":
            out.append(f"> ✅ **Test it:** {b[1]}\n")
        elif k == "rule":
            out.append("---\n")
    return "\n".join(out).strip() + "\n"

# ---- docx helpers ----
BRAND = RGBColor(0x1F,0x6F,0xEB); DARK = RGBColor(0x11,0x18,0x27); GREY = RGBColor(0x55,0x5B,0x66)
def _shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexc); tcPr.append(shd)

# inline **bold** + `code` rendering for docx
import re
def _runs(paragraph, text):
    # split on **bold** and `code`
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part: continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.color.rgb = RGBColor(0xC7,0x25,0x4E)
        else:
            paragraph.add_run(part)

def render_docx(blocks):
    import prodoc
    doc = Document()
    nrm = doc.styles["Normal"]; nrm.font.name = "Arial"; nrm.font.size = Pt(11)
    prodoc.style_headings(doc)
    # --- professional front matter ---
    prodoc.add_cover_page(doc, "Learner Guide", TITLE, VERSION,
        org_logo=os.path.join(REPO,"courseware/assets/tertiary-infotech-logo.png"),
        course_logo=os.path.join(REPO,"courseware/assets/n8n-course-logo.png"))
    prodoc.add_version_control(doc, VERSIONS)
    prodoc.add_toc(doc, levels="1-2")
    for b in blocks:
        k = b[0]
        if k == "h1":
            continue  # title is on the cover page
        elif k == "h2":
            pr = doc.add_paragraph(style="Heading 1"); pr.add_run(b[1])
        elif k == "h3":
            pr = doc.add_paragraph(style="Heading 2"); pr.add_run(b[1])
        elif k == "p":
            pr = doc.add_paragraph(); _runs(pr, b[1])
        elif k == "steps":
            for i, s in enumerate(b[1], 1):
                pr = doc.add_paragraph(style="List Number"); _runs(pr, s)
        elif k == "bullets":
            for s in b[1]:
                pr = doc.add_paragraph(style="List Bullet"); _runs(pr, s)
        elif k == "code":
            pr = doc.add_paragraph(); _shade_para(pr)
            r = pr.add_run(b[1]); r.font.name = "Consolas"; r.font.size = Pt(9)
        elif k == "table":
            rows = b[1]
            t = doc.add_table(rows=0, cols=len(rows[0])); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci, val in enumerate(row):
                    cells[ci].text = ""
                    pp = cells[ci].paragraphs[0]
                    if ri == 0:
                        rr = pp.add_run(val); rr.bold = True; rr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); rr.font.size = Pt(9.5)
                        _shade(cells[ci], "1F6FEB")
                    else:
                        _runs(pp, val)
                        for rn in pp.runs: rn.font.size = Pt(9.5)
        elif k == "note":
            pr = doc.add_paragraph(); _shade_para(pr, "FFF4E5")
            rr = pr.add_run("Note:  "); rr.bold = True; rr.font.color.rgb = RGBColor(0xB5,0x6A,0x00)
            _runs(pr, b[1])
        elif k == "test":
            pr = doc.add_paragraph(); _shade_para(pr, "E8F7EE")
            rr = pr.add_run("✅ Test it:  "); rr.bold = True; rr.font.color.rgb = RGBColor(0x12,0x7A,0x3E)
            _runs(pr, b[1])
        elif k == "img":
            try:
                pic = doc.add_paragraph(); pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(os.path.join(REPO, b[1]), width=Inches(6.2))
                cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cr = cap.add_run(b[2]); cr.italic = True; cr.font.size = Pt(9); cr.font.color.rgb = GREY
            except Exception as e:
                print("  [img skip]", b[1], e)
        elif k == "rule":
            pr = doc.add_paragraph(); pr.paragraph_format.space_before = Pt(2); pr.paragraph_format.space_after = Pt(2)
            ppr = pr._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6"); bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),"D0D7DE")
            bdr.append(bot); ppr.append(bdr)
    # page numbering + update-fields-on-open
    prodoc.add_page_numbers(doc)
    prodoc.enable_update_fields(doc)
    return doc

def _shade_para(pr, hexc="F3F5F8"):
    ppr = pr._p.get_or_add_pPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hexc); ppr.append(shd)

# ============================================================================
# WRITE BOTH
# ============================================================================
ACT_IMG = {
    "Activity 1 ": ("labs/activity1-flyer-form/Activity1-Flyer-Form.png", "Activity 1 workflow — Form Trigger to Gmail"),
    "Activity 2 ": ("labs/activity2-data-table/Activity2-Data-Table.png", "Activity 2 workflow — Form to Gmail + Data Table"),
    "Activity 3a": ("labs/activity3-conditional/Activity3a-Conditional-Data-Table.png", "Activity 3a workflow — IF routing to Data Table / email"),
    "Activity 3b": ("labs/activity3-conditional/Activity3b-Conditional-Google-Sheets.png", "Activity 3b workflow — IF routing to Google Sheets / email"),
    "Activity 4a": ("labs/activity4-telegram-agent/Activity4a-Telegram-Agent.png", "Activity 4a workflow — Telegram-triggered AI Agent"),
    "Activity 4b": ("labs/activity4-telegram-agent/Activity4b-Telegram-Data-Table.png", "Activity 4b workflow — Agent with a Data Table tool"),
    "Activity 5 ": ("labs/activity5-rag/Activity5-RAG-Telegram.png", "Activity 5 workflow — Agent with RAG + Data Table sources"),
    "Activity 6 ": ("labs/activity6-investment-advisor/Activity6-Investment-Advisor.png", "Activity 6 workflow — Webhook chatbot + enquiry"),
    "Activity 7 ": ("labs/activity7-finance-advisor/Activity7-Finance-Advisor.png", "Activity 7 workflow — Finance API to Telegram day trader"),
    "Activity 8a": ("labs/activity8-guardrails/Activity8a-Human-in-the-Loop.png", "Activity 8a workflow — human-in-the-loop approval"),
    "Activity 8b": ("labs/activity8-guardrails/Activity8b-Guardrails.png", "Activity 8b workflow — pre/post guardrails around the agent"),
}

def insert_images(blocks):
    """After each activity's 'What you'll build' / 'Workflow overview' flow line, insert its diagram."""
    out = []; cur = None; armed = False
    for b in blocks:
        out.append(b)
        if b[0] == "h2":
            cur = next((v for k, v in ACT_IMG.items() if b[1].startswith(k)), None); armed = False
        elif b[0] == "h3" and cur and b[1].strip().lower() == "goal":
            armed = True
        elif b[0] == "p" and armed and cur:
            out.append(("img", cur[0], cur[1])); armed = False; cur2 = cur; cur = None
    return out

B = insert_images(B)
md = render_markdown(B)
with open(os.path.join(REPO, "LEARNER-GUIDE.md"), "w") as f:
    f.write(md)
render_docx(B).save(os.path.join(REPO, "courseware/n8n Automation Learner Guide.docx"))

# alignment report
n_h2 = sum(1 for b in B if b[0]=="h2")
print("Wrote LEARNER-GUIDE.md and courseware DOCX from one source.")
print("Sections (h2):", n_h2, "| total blocks:", len(B), "| md chars:", len(md))
