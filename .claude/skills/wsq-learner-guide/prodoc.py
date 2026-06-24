"""Shared professional-document helpers matching the WSQ house format:
cover page, Document Version Control Record, Table of Contents (field),
"Page X of Y" footer, branded Heading styles, and update-fields-on-open."""
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x1F, 0x6F, 0xEB)
DARK  = RGBColor(0x11, 0x18, 0x27)
GREY  = RGBColor(0x55, 0x5B, 0x66)

ORG  = "Tertiary Infotech Academy Pte Ltd"
COPYRIGHT = "© 2026 Tertiary Infotech Academy Pte Ltd. All rights reserved."
UEN  = "UEN: 201200696W"
TGS  = "TGS Ref No: TGS-2023035977"

def _shade_cell(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexc); tcPr.append(shd)

def _field(paragraph, instr, default=""):
    """Insert a Word field (e.g. PAGE, NUMPAGES, TOC) into a paragraph."""
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = default
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(it); run._r.append(sep); run._r.append(t); run._r.append(end)
    return run

def _center(p):
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; return p

def _line(doc, text, size, bold=False, color=DARK, before=0, after=4, caps=False):
    p = doc.add_paragraph(); _center(p)
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text.upper() if caps else text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color
    return p

def style_headings(doc):
    """Brand the built-in Heading styles so the TOC field can capture them."""
    specs = {"Title": (24, DARK, True), "Heading 1": (16, BRAND, True),
             "Heading 2": (13, DARK, True), "Heading 3": (11, BRAND, True)}
    for name, (sz, col, bold) in specs.items():
        try:
            st = doc.styles[name]; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = bold
            st.font.name = "Arial"
        except KeyError:
            pass

def _img(doc, path, width_in):
    from docx.shared import Inches
    import os
    if path and os.path.exists(path):
        p = doc.add_paragraph(); _center(p)
        p.add_run().add_picture(path, width=Inches(width_in))
        return True
    return False

def add_cover_page(doc, kind, title, version, conducted_by=ORG,
                   org_logo=None, course_logo=None):
    """kind: 'Learner Guide' or 'LESSON PLAN'.
    org_logo: Tertiary Infotech Academy logo path; course_logo: n8n course logo path."""
    doc.add_paragraph()
    _img(doc, org_logo, 2.1)                              # Tertiary Infotech Academy logo
    _line(doc, ORG, 13, bold=True, color=DARK, before=2, after=1)
    _line(doc, UEN, 10, color=GREY, after=10)
    _line(doc, kind.upper(), 26, bold=True, color=BRAND, after=4)
    _line(doc, "For", 12, color=GREY, before=4, after=8)
    _img(doc, course_logo, 1.0)                           # n8n course logo
    _line(doc, title, 20, bold=True, color=DARK, before=4, after=2)   # Course Title
    _line(doc, TGS, 12, color=DARK, before=4, after=14)
    _line(doc, "Conducted by", 11, color=GREY, after=2)
    _line(doc, conducted_by, 13, bold=True, color=DARK, after=2)
    _line(doc, UEN, 11, color=GREY, after=12)
    _line(doc, f"Version {version}", 12, bold=True, color=BRAND)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_version_control(doc, rows):
    """rows: list of (version, effective_date, changes, author)."""
    h = doc.add_paragraph(); r = h.add_run("DOCUMENT VERSION CONTROL RECORD")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = DARK
    h.paragraph_format.space_after = Pt(6)
    t = doc.add_table(rows=0, cols=4); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ["Version Number", "Effective Date of Release", "Summary of Included Changes", "Author"]
    hr = t.add_row().cells
    for i, htext in enumerate(hdr):
        hr[i].text = ""; rr = hr[i].paragraphs[0].add_run(htext)
        rr.bold = True; rr.font.size = Pt(9.5); rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hr[i], "1F6FEB")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; rr = cells[i].paragraphs[0].add_run(str(val)); rr.font.size = Pt(9.5)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_toc(doc, levels="1-3"):
    h = doc.add_paragraph(); r = h.add_run("TABLE OF CONTENTS")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = DARK
    h.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    _field(p, f'TOC \\o "{levels}" \\h \\z \\u',
           default="Right-click and choose “Update Field”, or press F9, to build the table of contents.")
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def add_page_numbers(doc, left_text=""):
    sec = doc.sections[0]; footer = sec.footer; footer.is_linked_to_previous = False
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    r = p.add_run("Page "); r.font.size = Pt(9); r.font.color.rgb = GREY
    _field(p, "PAGE", "1").font.size = Pt(9)
    r2 = p.add_run(" of "); r2.font.size = Pt(9); r2.font.color.rgb = GREY
    _field(p, "NUMPAGES", "1").font.size = Pt(9)
    if left_text:
        hp = sec.footer.add_paragraph() if False else p  # keep single centered line
    # brand strip
    sp = footer.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(COPYRIGHT + "  ·  www.tertiarycourses.com.sg")
    sr.font.size = Pt(7.5); sr.font.color.rgb = GREY

def enable_update_fields(doc):
    """Tell Word to refresh all fields (TOC, PAGE) when the document opens."""
    settings = doc.settings.element
    uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true"); settings.append(uf)
