#!/usr/bin/env python3
"""Generate the 3-day Lesson Plan Word document for the WSQ course
'Agentic AI Automation with n8n' (TGS-2023035977) — revised flow."""

import os, sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import prodoc
REPO_DIR = "/Users/alfredang/projects/labs-activities/TGS-2023035977-Agentic AI Automation with n8n"
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

LP_VERSIONS = [
    ("1.0", "24 June 2026", "First version — 3-day lesson plan aligned to the agentic n8n course flow",
     "Tertiary Infotech Academy Pte Ltd"),
]

BRAND = RGBColor(0x1F, 0x6F, 0xEB)
DARK  = RGBColor(0x11, 0x18, 0x27)
GREY  = RGBColor(0x55, 0x5B, 0x66)
HEADER_FILL = "1F6FEB"
TOPIC_FILL  = "E8F0FE"
BREAK_FILL  = "FFF4E5"

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=10, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.italic = italic; run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p

def add_heading(text, size=15, color=BRAND, space_before=12, space_after=6):
    # Use the built-in Heading 1 style so the Table of Contents field captures it
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.add_run(text)
    return p

# ============================ COVER + FRONT MATTER ============================
prodoc.style_headings(doc)
prodoc.add_cover_page(doc, "LESSON PLAN", "Agentic AI Automation with n8n", "1.0",
    org_logo=os.path.join(REPO_DIR,"courseware/assets/tertiary-infotech-logo.png"),
    course_logo=os.path.join(REPO_DIR,"courseware/assets/n8n-course-logo.png"))
prodoc.add_version_control(doc, LP_VERSIONS)
prodoc.add_toc(doc, levels="1-2")

# ============================ COURSE INFO ============================
info = [
    ("Course Title", "Agentic AI Automation with n8n"),
    ("Course Code", "TGS-2023035977"),
    ("Duration", "3 Days (24 training hours)"),
    ("Daily Schedule", "9:30 AM – 6:30 PM (8 training hours/day, excluding lunch)"),
    ("Lunch Break", "1:00 PM – 2:00 PM (1 hour)"),
    ("Breaks", "Short tea breaks are scheduled within each day's training hours"),
    ("Delivery Mode", "Instructor-led, hands-on labs with group work and presentations"),
    ("Prerequisites", "Basic computer literacy; a Google/Gmail account; no coding experience required"),
    ("Tools & Accounts", "n8n (trial + local Docker), OpenAI & Gemini API keys, Telegram, "
                         "Gmail/Outlook, Google Sheets, Twelve Data & NewsAPI, Claude Code"),
]
tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in info:
    row = tbl.add_row().cells
    set_cell_text(row[0], k, bold=True, color=DARK, size=10); set_cell_bg(row[0], "F1F5FB")
    set_cell_text(row[1], v, size=10)
for row in tbl.rows:
    row.cells[0].width = Inches(1.9); row.cells[1].width = Inches(4.6)

# ============================ OVERVIEW ============================
add_heading("Course Overview")
doc.add_paragraph(
    "This 3-day hands-on course teaches participants to design, build and secure agentic AI "
    "automations with n8n. Learners set up n8n (trial and local Docker), automate forms and email, "
    "store data in Data Tables and external spreadsheets, build Telegram-triggered AI agents, add "
    "Retrieval-Augmented Generation (RAG), integrate webhooks and external APIs, and apply "
    "human-in-the-loop guardrails — culminating in a mini capstone project that is presented and assessed."
)

# ============================ LEARNING OUTCOMES ============================
add_heading("Learning Outcomes")
doc.add_paragraph("By the end of this course, participants will be able to:")
for o in [
    "Set up n8n using a cloud trial and a local Docker Compose installation, and navigate the n8n UI.",
    "Apply the core concepts of triggers, actions, nodes and flows to build automated workflows.",
    "Automate form submissions with QR codes, send admin emails and capture data in Data Tables and Google Sheets/Excel.",
    "Design Telegram-triggered AI agents using LLMs, memory, tools and system instructions.",
    "Implement Retrieval-Augmented Generation (RAG) and route an agent between multiple data sources.",
    "Expose workflows via webhooks and integrate external services through APIs and HTTP requests.",
    "Apply human-in-the-loop and guardrail patterns to make AI automations safe and reliable.",
    "Plan, build, present and self-assess an end-to-end mini capstone automation project.",
]:
    p = doc.add_paragraph(o, style="List Bullet"); p.paragraph_format.space_after = Pt(2)

# ============================ SCHEDULE HELPERS ============================
def schedule_table(rows):
    table = doc.add_table(rows=1, cols=3); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Time", "Topic / Activity", "Duration"]):
        set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=10,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(hdr[i], HEADER_FILL)
    for time, activity, minutes, kind in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], time, size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, bold=(kind=="topic"))
        set_cell_text(cells[1], activity, size=9.5, bold=(kind in ("topic","break")),
                      color=BRAND if kind=="topic" else (GREY if kind=="break" else None),
                      italic=(kind=="break"))
        set_cell_text(cells[2], (f"{minutes} min" if minutes else "—"), size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill = TOPIC_FILL if kind=="topic" else (BREAK_FILL if kind=="break" else None)
        if fill:
            for cc in cells: set_cell_bg(cc, fill)
    widths = [Inches(1.2), Inches(4.45), Inches(0.85)]
    for row in table.rows:
        for i, w in enumerate(widths): row.cells[i].width = w
    return table

def day_header(day, theme):
    add_heading(day, size=14, color=DARK, space_before=16)
    p = doc.add_paragraph(); r = p.add_run(theme); r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(4)

def session_label(text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = BRAND

# ----------------------------- DAY 1 -----------------------------
day_header("Day 1", "Workflow Automation with n8n  ·  AI Agents & RAG")

session_label("Morning Session  ·  9:30 AM – 1:00 PM")
schedule_table([
    ("9:30 – 9:45",  "Welcome, course objectives & overview (TGS-2023035977)", 15, "normal"),
    ("9:45 – 12:25", "Topic 1: Workflow Automation with n8n", None, "topic"),
    ("9:45 – 10:05", "Overview of n8n", 20, "normal"),
    ("10:05 – 10:35","Set up n8n — create trial accounts (15 learners) and install locally with Docker Compose", 30, "normal"),
    ("10:35 – 10:50","Explore the n8n UI", 15, "normal"),
    ("10:50 – 11:10","Key concepts of n8n — Triggers, Actions, Nodes, Flows", 20, "normal"),
    ("11:10 – 11:35","Activity 1: QR Code form (name, email, phone, message) + email to admin via Gmail/Outlook (2 nodes)", 25, "normal"),
    ("11:35 – 12:25","Group Activity (3–4 pax): design an event flyer (e.g. bowling), generate a QR code with the QR generator, build the ad, then present (review past-student examples)", 50, "normal"),
    ("12:25 – 1:00", "Activity 2: Add an n8n Data Table to Activity 1 to capture submissions (3 nodes)", 35, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])

session_label("Afternoon Session  ·  2:00 PM – 6:30 PM")
schedule_table([
    ("2:00 – 2:40",  "Activity 3a: Conditional response — if yes, save the date to the Data Table; if no, send a thank-you email (note: trial Data Table data is not persistent)", 40, "normal"),
    ("2:40 – 3:20",  "Activity 3b: Conditional response — replace the Data Table with Google Sheet/Excel for persistent storage (Google training accounts); if no, send a thank-you email", 40, "normal"),
    ("3:20 – 3:35",  "Tea Break", 15, "break"),
    ("3:35 – 6:05",  "Topic 2: AI Agents and RAG", None, "topic"),
    ("3:35 – 4:05",  "Overview of AI Agents — LLM, Memory, Tools and System Instruction", 30, "normal"),
    ("4:05 – 4:50",  "Activity 4a: Telegram-triggered AI Agent — simple customer-service chatbot (OpenAI / Gemini API key provided)", 45, "normal"),
    ("4:50 – 6:05",  "Activity 4b: Telegram-triggered AI Agent (customer service / HR admin) — attach a Data Table as a tool; reuse Activity 2 data or upload a 100-record employee CSV", 75, "normal"),
    ("6:05 – 6:30",  "Day 1 recap, Q&A and wrap-up", 25, "normal"),
])

# ----------------------------- DAY 2 -----------------------------
day_header("Day 2", "RAG  ·  Webhooks  ·  APIs & HTTP Requests")

session_label("Morning Session  ·  9:30 AM – 1:00 PM")
schedule_table([
    ("9:30 – 9:45",  "Day 1 recap & Day 2 objectives", 15, "normal"),
    ("9:45 – 1:00",  "Topic 2 (cont.): Retrieval-Augmented Generation (RAG)", None, "topic"),
    ("9:45 – 10:30", "Overview of RAG — Tokenization, Embeddings, Vector Stores", 45, "normal"),
    ("10:30 – 11:00","Generate mock documents with Claude Code (e.g. employee-benefits FAQ, product info)", 30, "normal"),
    ("11:00 – 11:15","Tea Break", 15, "break"),
    ("11:15 – 12:30","Activity 5: Add RAG to the Telegram chatbot — route between two data sources (RAG vector store + the earlier Data Table) via the system instruction", 75, "normal"),
    ("12:30 – 1:00", "Selected learners present their chatbots", 30, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])

session_label("Afternoon Session  ·  2:00 PM – 6:30 PM")
schedule_table([
    ("2:00 – 4:00",  "Topic 3: Webhooks", None, "topic"),
    ("2:00 – 2:30",  "Overview of Webhooks — use cases and external triggers", 30, "normal"),
    ("2:30 – 3:40",  "Activity 6: Website chatbot via webhook — follow the n8n Investment Advisor reference site", 70, "normal"),
    ("3:40 – 4:00",  "Selected learners present their website and chatbot", 20, "normal"),
    ("4:00 – 4:15",  "Tea Break", 15, "break"),
    ("4:15 – 6:30",  "Topic 4: API and HTTP Request", None, "topic"),
    ("4:15 – 4:45",  "Overview of APIs and HTTP Requests", 30, "normal"),
    ("4:45 – 6:05",  "Activity 7: Pull data from a Finance API (Twelve Data) and news (NewsAPI) to reply on Telegram — learners obtain their own API keys", 80, "normal"),
    ("6:05 – 6:30",  "Day 2 recap, Q&A and wrap-up", 25, "normal"),
])

# ----------------------------- DAY 3 -----------------------------
day_header("Day 3", "Security & Guardrails  ·  Mini Capstone Project")

session_label("Morning Session  ·  9:30 AM – 1:00 PM")
schedule_table([
    ("9:30 – 9:45",  "Day 2 recap & Day 3 objectives", 15, "normal"),
    ("9:45 – 1:00",  "Topic 5: Security and Guardrails", None, "topic"),
    ("9:45 – 10:25", "Overview of Human-in-the-Loop", 40, "normal"),
    ("10:25 – 11:25","Activity 8a: Build an approval flow with human-in-the-loop — leave application approval", 60, "normal"),
    ("11:25 – 11:40","Tea Break", 15, "break"),
    ("11:40 – 12:00","Selected learners present their approval flow", 20, "normal"),
    ("12:00 – 12:25","Overview of Guardrails", 25, "normal"),
    ("12:25 – 1:00", "Activity 8b: Add pre- and post-guardrails in front of the AI agent", 35, "normal"),
])
schedule_table([("1:00 – 2:00", "Lunch Break", 60, "break")])

session_label("Afternoon Session  ·  2:00 PM – 6:30 PM")
schedule_table([
    ("2:00 – 6:30",  "Topic 6: Mini Capstone Project", None, "topic"),
    ("2:00 – 2:30",  "Capstone briefing — requirements, scope and assessment criteria", 30, "normal"),
    ("2:30 – 4:15",  "Mini capstone build (supervised lab time)", 105, "normal"),
    ("4:15 – 4:30",  "Tea Break", 15, "break"),
    ("4:30 – 5:45",  "Presentation of mini capstone projects", 75, "normal"),
    ("5:45 – 6:30",  "Assessment and course closing", 45, "normal"),
])

# ============================ RESOURCES ============================
add_heading("Tools, Accounts & Resources", space_before=16)
res = [
    ("QR Code Generator", "https://alfredang.github.io/qrcodegenerator/"),
    ("Webhook reference — n8n Investment Advisor", "https://alfredang.github.io/n8n-investmentadvisor/"),
    ("Finance API — Twelve Data", "https://twelvedata.com/login"),
    ("News API — NewsAPI", "https://newsapi.org/"),
    ("LLM API keys", "OpenAI API key & Google Gemini API key (provided to learners)"),
    ("Local n8n", "Docker Compose self-hosted install (see Learner Guide)"),
]
rt = doc.add_table(rows=0, cols=2); rt.style = "Table Grid"; rt.alignment = WD_TABLE_ALIGNMENT.CENTER
for k, v in res:
    row = rt.add_row().cells
    set_cell_text(row[0], k, bold=True, color=DARK, size=9.5); set_cell_bg(row[0], "F1F5FB")
    set_cell_text(row[1], v, size=9.5, color=BRAND)
for row in rt.rows:
    row.cells[0].width = Inches(2.4); row.cells[1].width = Inches(4.1)

# ============================ ASSESSMENT ============================
add_heading("Assessment", space_before=14)
doc.add_paragraph(
    "Participants are assessed through their hands-on activities across the three days and a mini "
    "capstone project on Day 3. The capstone is presented to the class and evaluated against the "
    "course learning outcomes — workflow design, AI agent and RAG integration, webhook and API "
    "use, and the application of security guardrails."
)

# ============================ FOOTER (page numbers) ============================
prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)

OUT = "courseware/Lesson Plan - Agentic AI Automation with n8n.docx"
doc.save(OUT)
print("Saved:", OUT)

# ---- sanity check: every day = 480 min (incl. tea breaks, excl. lunch) ----
days = {
    "Day 1": [15,20,30,15,20,25,50,35] + [40,40,15,30,45,75,25],
    "Day 2": [15,45,30,15,75,30] + [30,70,20,15,30,80,25],
    "Day 3": [15,40,60,15,20,25,35] + [30,105,15,75,45],
}
for d, mins in days.items():
    total = sum(mins)
    print(f"{d}: {total} min = {total/60:.2f} hrs ->", "OK" if total == 480 else "MISMATCH")
